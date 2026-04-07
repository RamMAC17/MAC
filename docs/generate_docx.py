"""Generate professional DOCX documents for each phase of the MAC (MBM AI Cloud) project."""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
from pathlib import Path

OUT = Path(__file__).parent
plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['font.sans-serif'] = ['Segoe UI', 'Arial', 'Helvetica', 'DejaVu Sans']

# ─── Colour palette ───
C_DARK    = '#0d1b2a'
C_PRIMARY = '#1b4965'
C_SEC     = '#2a6f97'
C_ACCENT  = '#468faf'
C_LIGHT   = '#a9d6e5'
C_PALE    = '#e8f4f8'
C_ORANGE  = '#e76f51'
C_GREEN   = '#52b788'
C_TEAL    = '#2ec4b2'
C_AMBER   = '#f4a261'
C_RED     = '#e63946'
C_GRAY    = '#6c757d'
C_LGRAY   = '#f6f8fa'

# ════════════════════════════════════════════════════════════════
#  HELPERS
# ════════════════════════════════════════════════════════════════

def _shade(cell, color_hex):
    """Set background colour on a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex.lstrip('#'))
    shading.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading)

def _set_col_widths(table, widths_inches):
    for row in table.rows:
        for i, w in enumerate(widths_inches):
            row.cells[i].width = Inches(w)

def new_doc(title, subtitle, date='07 April 2026'):
    """Create a new Document with a styled title page."""
    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.54)
        sec.right_margin = Cm(2.54)

    # Title page
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('MAC')
    r.font.size = Pt(44)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1b, 0x49, 0x65)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('MBM AI Cloud')
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x2a, 0x6f, 0x97)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(16)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(date)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)

    doc.add_page_break()
    return doc


def body(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header row
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        _shade(c, C_PRIMARY)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if ri % 2 == 1:
                _shade(c, C_LGRAY)

    if col_widths:
        _set_col_widths(t, col_widths)

    doc.add_paragraph()  # spacer
    return t


def add_diagram(doc, buf, width=6.2):
    """Insert a diagram image from a BytesIO buffer."""
    doc.add_picture(buf, width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def save(doc, name):
    path = OUT / name
    doc.save(str(path))
    print(f"  OK: {name} ({path.stat().st_size // 1024} KB)")


# ════════════════════════════════════════════════════════════════
#  DIAGRAM PRIMITIVES
# ════════════════════════════════════════════════════════════════

def _box(ax, x, y, w, h, label, color=C_PRIMARY, tc='white', fs=9, fw='bold', ls=None):
    b = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.03",
                       facecolor=color, edgecolor='#dddddd', linewidth=0.6)
    ax.add_patch(b)
    lines = label.split('\n')
    if len(lines) == 1:
        ax.text(x + w/2, y + h/2, label, ha='center', va='center',
                fontsize=fs, color=tc, fontweight=fw, fontstyle='italic' if ls else 'normal')
    else:
        top = lines[0]
        bot = '\n'.join(lines[1:])
        ax.text(x + w/2, y + h*0.62, top, ha='center', va='center',
                fontsize=fs, color=tc, fontweight=fw)
        ax.text(x + w/2, y + h*0.30, bot, ha='center', va='center',
                fontsize=fs - 1.5, color=tc, fontweight='normal', alpha=0.85)


def _arrow(ax, x1, y1, x2, y2, color='#555555', lw=1.5, style='->'):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle=style, color=color, lw=lw))


def _fig(w=10, h=6, xl=(0, 10), yl=(0, 6)):
    fig, ax = plt.subplots(figsize=(w, h))
    ax.set_xlim(xl)
    ax.set_ylim(yl)
    ax.axis('off')
    fig.patch.set_facecolor('white')
    return fig, ax


def _to_buf(fig):
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    return buf


# ════════════════════════════════════════════════════════════════
#  DIAGRAMS
# ════════════════════════════════════════════════════════════════

def dia_roadmap():
    """8-phase project roadmap."""
    fig, ax = _fig(12, 5.5, (0, 12), (0, 5.5))

    phases = [
        ('Phase 1', 'API Endpoints'),
        ('Phase 2', 'LLM Models'),
        ('Phase 3', 'API–Model\nIntegration'),
        ('Phase 4', 'Usage Control'),
        ('Phase 5', 'Web Interface'),
        ('Phase 6', 'Guardrails'),
        ('Phase 7', 'Knowledge\nBase + RAG'),
        ('Phase 8', 'Retrieval\n+ Search'),
    ]
    colors = ['#0d1b2a', '#1b3a4b', '#1b4965', '#2a6f97',
              '#468faf', '#61a5c2', '#7fb8cc', '#52b788']
    tcs = ['white'] * 5 + ['white', 'white', 'white']

    xs = [0.3, 3.1, 5.9, 8.7]
    bw, bh = 2.4, 1.1

    for i, (label, sub) in enumerate(phases):
        row = 0 if i < 4 else 1
        col = i % 4
        x = xs[col]
        y = 3.5 if row == 0 else 1.5
        _box(ax, x, y, bw, bh, f'{label}\n{sub}', colors[i], tcs[i], fs=10)

    # Arrows row 1
    for i in range(3):
        _arrow(ax, xs[i] + bw, 3.5 + bh/2, xs[i+1], 3.5 + bh/2)
    # Curved arrow Phase 4 → Phase 5
    ax.annotate('', xy=(xs[0] + bw/2, 1.5 + bh), xytext=(xs[3] + bw/2, 3.5),
                arrowprops=dict(arrowstyle='->', color='#555', lw=1.5,
                                connectionstyle='arc3,rad=0.4'))
    # Arrows row 2
    for i in range(3):
        _arrow(ax, xs[i] + bw, 1.5 + bh/2, xs[i+1], 1.5 + bh/2)

    ax.text(6, 5.1, 'MAC — Build Roadmap', ha='center', fontsize=14,
            fontweight='bold', color=C_DARK)
    return _to_buf(fig)


def dia_architecture():
    """Full system architecture diagram."""
    fig, ax = _fig(11, 8.5, (0, 11), (0, 8.5))

    # Title
    ax.text(5.5, 8.1, 'System Architecture', ha='center', fontsize=14,
            fontweight='bold', color=C_DARK)

    # Layer: Users
    _box(ax, 2.5, 7, 6, 0.7, 'Students / Faculty — College LAN', C_TEAL, 'white', 10)

    _arrow(ax, 5.5, 7, 5.5, 6.5)

    # Layer: Nginx
    _box(ax, 3, 5.7, 5, 0.7, 'Nginx — Reverse Proxy · TLS · Caching', C_ORANGE, 'white', 10)

    _arrow(ax, 5.5, 5.7, 5.5, 5.2)

    # Layer: FastAPI
    _box(ax, 1.8, 4.4, 7.4, 0.7, 'FastAPI — API Gateway · Authentication · Rate Limiting · Routing', C_PRIMARY, 'white', 10)

    # Arrows down to services
    _arrow(ax, 2.5, 4.4, 1.5, 3.7)
    _arrow(ax, 5.5, 4.4, 5.5, 3.7)
    _arrow(ax, 8.5, 4.4, 9.5, 3.7)

    # Layer: Services
    _box(ax, 0.2, 2.9, 2.6, 0.7, 'PostgreSQL\nUsers · Logs · Keys', C_SEC, 'white', 9)
    _box(ax, 4, 2.9, 3, 0.7, 'LiteLLM Proxy\nRouting · Load Balance', C_SEC, 'white', 9)
    _box(ax, 8.2, 2.9, 2.6, 0.7, 'Redis · Qdrant\nCache · Vectors', C_SEC, 'white', 9)

    # Arrow from LiteLLM to vLLM
    _arrow(ax, 5.5, 2.9, 5.5, 2.4)

    # Layer: vLLM
    _box(ax, 3, 1.6, 5, 0.7, 'vLLM — Model Workers · GPU Inference', C_DARK, 'white', 10)

    # GPU icon text
    _box(ax, 1.5, 0.4, 2.2, 0.6, 'GPU Node 1', C_ACCENT, 'white', 8)
    _box(ax, 4.4, 0.4, 2.2, 0.6, 'GPU Node 2', C_ACCENT, 'white', 8)
    _box(ax, 7.3, 0.4, 2.2, 0.6, 'GPU Node N', C_ACCENT, 'white', 8)
    ax.text(6.85, 0.65, '...', fontsize=16, color=C_GRAY, ha='center')
    _arrow(ax, 4, 1.6, 2.6, 1.05)
    _arrow(ax, 5.5, 1.6, 5.5, 1.05)
    _arrow(ax, 7, 1.6, 8.4, 1.05)

    return _to_buf(fig)


def dia_auth_flow():
    """Authentication flow diagram."""
    fig, ax = _fig(12, 4, (0, 12), (0, 4))
    ax.text(6, 3.7, 'Authentication Flow', ha='center', fontsize=13, fontweight='bold', color=C_DARK)

    steps = [
        (0.2, 'Student\nClient', C_TEAL),
        (2.4, 'POST\n/auth/login', C_ACCENT),
        (4.6, 'Validate\nCredentials', C_SEC),
        (6.8, 'Generate\nJWT + Refresh', C_PRIMARY),
        (9.0, 'Return\nTokens + Profile', C_GREEN),
    ]
    bw, bh = 2, 1.2
    y = 1.3
    for x, label, color in steps:
        _box(ax, x, y, bw, bh, label, color, 'white', 9)
    for i in range(len(steps) - 1):
        _arrow(ax, steps[i][0] + bw, y + bh/2, steps[i+1][0], y + bh/2)

    # Below: refresh flow
    ax.text(1.2, 0.6, 'Token expired?', fontsize=8, color=C_GRAY, ha='center', fontstyle='italic')
    _arrow(ax, 1.2, 0.8, 3.4, 1.3, color=C_AMBER, style='->', lw=1.0)
    ax.text(3.4, 0.6, 'POST /auth/refresh → new access token', fontsize=8, color=C_AMBER, ha='left')

    return _to_buf(fig)


def dia_model_vram():
    """Model VRAM requirements bar chart."""
    fig, ax = plt.subplots(figsize=(9, 4.5))
    models = ['Whisper\nLarge v3', 'Qwen2.5\nCoder 7B', 'DeepSeek\nR1 8B', 'LLaVA 1.6\n7B', 'Qwen2.5\n14B']
    vram = [3, 5, 6, 8, 10]
    colors = [C_ACCENT, C_SEC, C_PRIMARY, C_ORANGE, C_DARK]

    bars = ax.barh(models, vram, color=colors, height=0.55, edgecolor='white', linewidth=0.5)
    for bar, v in zip(bars, vram):
        ax.text(bar.get_width() + 0.15, bar.get_y() + bar.get_height()/2,
                f'{v} GB', va='center', fontsize=10, fontweight='bold', color=C_DARK)

    ax.set_xlabel('VRAM Required (GB)', fontsize=10, color=C_DARK)
    ax.set_title('Model VRAM Requirements', fontsize=13, fontweight='bold', color=C_DARK, pad=12)
    ax.set_xlim(0, 13)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.tick_params(left=False, labelsize=9)
    ax.grid(axis='x', alpha=0.2)
    fig.tight_layout()
    return _to_buf(fig)


def dia_smart_routing():
    """Smart model routing decision flow."""
    fig, ax = _fig(12, 6.5, (0, 12), (0, 6.5))
    ax.text(6, 6.2, 'Smart Routing — model: "auto"', ha='center', fontsize=13,
            fontweight='bold', color=C_DARK)

    # Input
    _box(ax, 4.5, 5, 3, 0.7, 'Incoming Request', C_TEAL, 'white', 10)
    _arrow(ax, 6, 5, 6, 4.5)

    # Analyser
    _box(ax, 4, 3.7, 4, 0.7, 'Content Analyser\nKeywords · Attachments', C_PRIMARY, 'white', 9)

    # Decision branches
    checks = [
        (0.3, 'Code\nkeywords?', 'Qwen2.5-Coder', C_SEC),
        (2.7, 'Math /\nreasoning?', 'DeepSeek-R1', C_ACCENT),
        (5.1, 'Image\nattached?', 'LLaVA 1.6', C_ORANGE),
        (7.5, 'Audio\nupload?', 'Whisper v3', C_AMBER),
        (9.9, 'General\ntext?', 'Qwen2.5-14B', C_GREEN),
    ]
    # arrows from analyser to decisions
    for x, q, m, col in checks:
        _arrow(ax, 6, 3.7, x + 1, 3.2)
        # Decision diamond (simulated with rotated box)
        _box(ax, x, 2.2, 2, 0.7, q, '#f0f0f0', C_DARK, 8, 'normal')
        _arrow(ax, x + 1, 2.2, x + 1, 1.6)
        _box(ax, x, 0.7, 2, 0.7, m, col, 'white', 8)

    return _to_buf(fig)


def dia_integration_arch():
    """API–Model integration architecture."""
    fig, ax = _fig(11, 6, (0, 11), (0, 6))
    ax.text(5.5, 5.7, 'Integration Architecture', ha='center', fontsize=13,
            fontweight='bold', color=C_DARK)

    _box(ax, 3.5, 4.6, 4, 0.7, 'FastAPI Gateway\n/query endpoints', C_PRIMARY, 'white', 10)
    _arrow(ax, 5.5, 4.6, 5.5, 4.1)

    _box(ax, 3, 3.3, 5, 0.7, 'LiteLLM Proxy\nModel Router · Load Balancer · Retry', C_SEC, 'white', 9)

    _arrow(ax, 3.5, 3.3, 1.5, 2.7)
    _arrow(ax, 5.5, 3.3, 5.5, 2.7)
    _arrow(ax, 7.5, 3.3, 9.5, 2.7)

    _box(ax, 0.2, 1.9, 2.6, 0.7, 'vLLM Worker 1\nqwen2.5-coder', C_DARK, 'white', 8)
    _box(ax, 4.2, 1.9, 2.6, 0.7, 'vLLM Worker 2\ndeepseek-r1', C_DARK, 'white', 8)
    _box(ax, 8.2, 1.9, 2.6, 0.7, 'vLLM Worker N\nqwen2.5-14b', C_DARK, 'white', 8)
    ax.text(7.4, 2.2, '...', fontsize=18, color=C_GRAY, ha='center')

    # Health checks
    _box(ax, 3.5, 0.6, 4, 0.6, 'Health Monitor\nLatency · GPU Temp · Queue', C_AMBER, 'white', 8)
    _arrow(ax, 1.5, 1.9, 4, 1.25, color=C_AMBER, lw=1, style='->')
    _arrow(ax, 5.5, 1.9, 5.5, 1.25, color=C_AMBER, lw=1, style='->')
    _arrow(ax, 9.5, 1.9, 7, 1.25, color=C_AMBER, lw=1, style='->')

    return _to_buf(fig)


def dia_rate_limit():
    """Rate limiting flow."""
    fig, ax = _fig(12, 4, (0, 12), (0, 4))
    ax.text(6, 3.7, 'Request Processing Pipeline', ha='center', fontsize=13,
            fontweight='bold', color=C_DARK)

    steps = [
        (0.1, 'Incoming\nRequest', C_TEAL),
        (2.2, 'Validate\nAPI Key', C_SEC),
        (4.3, 'Check\nRate Limit', C_ACCENT),
        (6.4, 'Check\nToken Quota', C_PRIMARY),
        (8.5, 'Route to\nModel', C_GREEN),
        (10.3, 'Return\nResponse', C_DARK),
    ]
    bw, bh = 1.7, 1.1
    y = 1.4
    for x, label, color in steps:
        _box(ax, x, y, bw, bh, label, color, 'white', 8)
    for i in range(len(steps) - 1):
        _arrow(ax, steps[i][0] + bw, y + bh/2, steps[i+1][0], y + bh/2)

    # Rejection arrows
    for i, label in [(1, '401'), (2, '429'), (3, '429')]:
        ax.annotate('', xy=(steps[i][0] + bw/2, y - 0.15),
                    xytext=(steps[i][0] + bw/2, y),
                    arrowprops=dict(arrowstyle='->', color=C_RED, lw=1.2))
        ax.text(steps[i][0] + bw/2, y - 0.4, f'Reject {label}', ha='center',
                fontsize=7, color=C_RED, fontweight='bold')

    return _to_buf(fig)


def dia_key_lifecycle():
    """API key lifecycle."""
    fig, ax = _fig(10, 5, (0, 10), (0, 5))
    ax.text(5, 4.7, 'API Key Lifecycle', ha='center', fontsize=13, fontweight='bold', color=C_DARK)

    _box(ax, 0.3, 2.5, 2, 0.8, 'Key\nGenerated', C_GREEN, 'white', 9)
    _arrow(ax, 2.3, 2.9, 3, 2.9)
    _box(ax, 3, 2.5, 2, 0.8, 'Active\n(In Use)', C_PRIMARY, 'white', 9)

    # Three outcomes
    _arrow(ax, 5, 3.1, 6.5, 3.8)
    _box(ax, 6.5, 3.5, 2.2, 0.7, 'Refreshed\n(Auto 30d)', C_ACCENT, 'white', 8)
    _arrow(ax, 8.7, 3.85, 9.3, 3.85, color=C_ACCENT)
    ax.annotate('', xy=(3, 3.3), xytext=(9.3, 3.85),
                arrowprops=dict(arrowstyle='->', color=C_ACCENT, lw=1,
                                connectionstyle='arc3,rad=0.5'))

    _arrow(ax, 5, 2.9, 6.5, 2.9)
    _box(ax, 6.5, 2.5, 2.2, 0.7, 'Manually\nRotated', C_AMBER, 'white', 8)

    _arrow(ax, 5, 2.7, 6.5, 1.8)
    _box(ax, 6.5, 1.4, 2.2, 0.7, 'Revoked /\nExpired', C_RED, 'white', 8)

    ax.text(1.3, 1.7, 'Static key: never expires\nRefresh key: auto-rotate 30 days',
            fontsize=8, color=C_GRAY, ha='center', linespacing=1.5)

    return _to_buf(fig)


def dia_dashboard():
    """Dashboard wireframe layout."""
    fig, ax = _fig(10, 7, (0, 10), (0, 7))
    ax.text(5, 6.7, 'Student Dashboard — Layout', ha='center', fontsize=13,
            fontweight='bold', color=C_DARK)

    # Header bar
    _box(ax, 0.2, 5.8, 9.6, 0.6, 'MAC — MBM AI Cloud                                                          Profile ▼', '#e8e8e8', C_DARK, 9)

    # Stat cards
    _box(ax, 0.3, 4.8, 2.8, 0.8, 'Tokens Today\n12,450', C_PRIMARY, 'white', 9)
    _box(ax, 3.5, 4.8, 2.8, 0.8, 'Requests Today\n23', C_SEC, 'white', 9)
    _box(ax, 6.7, 4.8, 3, 0.8, 'Quota Remaining\n75.1%', C_GREEN, 'white', 9)

    # Models + Quick Start
    _box(ax, 0.3, 2.3, 4.2, 2.2, '', '#f8f9fa', C_DARK, 9)
    ax.text(2.4, 4.2, 'Available Models', fontsize=9, fontweight='bold', color=C_DARK, ha='center')
    models_text = '[OK] Qwen2.5-Coder 7B\n[OK] DeepSeek-R1 8B\n[OK] Qwen2.5 14B\n[..] LLaVA (loading)\n[OK] Whisper v3'
    ax.text(0.6, 3.0, models_text, fontsize=7.5, color=C_DARK, va='center', linespacing=1.6, family='monospace')

    _box(ax, 5, 2.3, 4.7, 2.2, '', '#f8f9fa', C_DARK, 9)
    ax.text(7.3, 4.2, 'Quick Start', fontsize=9, fontweight='bold', color=C_DARK, ha='center')
    code_text = 'from openai import OpenAI\nclient = OpenAI(\n  base_url="http://mac/v1",\n  api_key="mac_sk_..."\n)\nresp = client.chat...'
    ax.text(5.3, 3.0, code_text, fontsize=7, color='#333', va='center', linespacing=1.5, family='monospace')

    # Usage chart
    _box(ax, 0.3, 0.3, 9.4, 1.7, '', '#f8f9fa', C_DARK, 9)
    ax.text(5, 1.8, 'Weekly Usage Trend', fontsize=9, fontweight='bold', color=C_DARK, ha='center')
    # Mini bar chart
    days = np.array([0.8, 2.0, 3.2, 4.4, 5.6, 6.8, 8.0])
    vals = np.array([0.6, 0.9, 1.2, 0.8, 1.1, 0.7, 0.4])
    for d, v in zip(days, vals):
        _box(ax, d, 0.45, 0.6, v, '', C_ACCENT, C_DARK, 1)

    return _to_buf(fig)


def dia_guardrails():
    """Guardrails pipeline diagram."""
    fig, ax = _fig(12, 5, (0, 12), (0, 5))
    ax.text(6, 4.7, 'Guardrails — Input & Output Filtering Pipeline', ha='center',
            fontsize=13, fontweight='bold', color=C_DARK)

    y = 2.0
    bw, bh = 1.8, 0.9
    pipeline = [
        (0.1, 'User\nInput', C_TEAL),
        (2.1, 'Prompt\nInjection\nDetection', C_RED),
        (4.1, 'Content\nFilter', C_ORANGE),
        (6.1, 'LLM\nInference', C_PRIMARY),
        (8.1, 'Output\nSafety\nFilter', C_AMBER),
        (10.1, 'Safe\nResponse', C_GREEN),
    ]
    for x, label, color in pipeline:
        _box(ax, x, y, bw, bh, label, color, 'white', 8)
    for i in range(len(pipeline) - 1):
        _arrow(ax, pipeline[i][0] + bw, y + bh/2, pipeline[i+1][0], y + bh/2)

    # Labels below
    checks_in = ['Jailbreak attempts', 'Blocked topics', 'Max prompt length']
    for i, txt in enumerate(checks_in):
        ax.text(3.1, 1.5 - i*0.35, f'• {txt}', fontsize=7.5, color=C_DARK)
    ax.text(3.1, 1.5 - 3*0.35 + 0.1, 'INPUT CHECKS', fontsize=7, color=C_RED, fontweight='bold')

    checks_out = ['PII redaction', 'Harmful content', 'Academic integrity notice']
    for i, txt in enumerate(checks_out):
        ax.text(8.1, 1.5 - i*0.35, f'• {txt}', fontsize=7.5, color=C_DARK)
    ax.text(8.1, 1.5 - 3*0.35 + 0.1, 'OUTPUT CHECKS', fontsize=7, color=C_AMBER, fontweight='bold')

    return _to_buf(fig)


def dia_rag_pipeline():
    """RAG pipeline — ingestion + query."""
    fig, ax = _fig(12, 7, (0, 12), (0, 7))
    ax.text(6, 6.7, 'RAG Pipeline — Ingestion & Query', ha='center',
            fontsize=13, fontweight='bold', color=C_DARK)

    # Ingestion flow (top)
    ax.text(1, 5.9, 'Document Ingestion', fontsize=10, fontweight='bold', color=C_SEC)
    ing = [
        (0.2, 'Upload\nPDF/DOCX', C_TEAL),
        (2.5, 'Chunking\n512 tokens', C_SEC),
        (4.8, 'Embedding\n768-dim', C_ACCENT),
        (7.1, 'Store in\nQdrant', C_PRIMARY),
    ]
    bw, bh = 2, 0.8
    y = 4.8
    for x, label, color in ing:
        _box(ax, x, y, bw, bh, label, color, 'white', 9)
    for i in range(len(ing) - 1):
        _arrow(ax, ing[i][0] + bw, y + bh/2, ing[i+1][0], y + bh/2)

    # Divider
    ax.plot([0.5, 11.5], [4.3, 4.3], color='#ddd', lw=1, ls='--')

    # Query flow (bottom)
    ax.text(1, 4.0, 'Query & Retrieval', fontsize=10, fontweight='bold', color=C_SEC)
    q_top = [
        (0.2, 'Student\nQuestion', C_TEAL),
        (2.5, 'Embed\nQuery', C_SEC),
        (4.8, 'Similarity\nSearch (top-k)', C_ACCENT),
    ]
    y2 = 3.0
    for x, label, color in q_top:
        _box(ax, x, y2, bw, bh, label, color, 'white', 9)
    for i in range(len(q_top) - 1):
        _arrow(ax, q_top[i][0] + bw, y2 + bh/2, q_top[i+1][0], y2 + bh/2)

    # Retrieved chunks + LLM
    _arrow(ax, 6.8, y2 + bh/2, 7.1, y2 + bh/2)
    _box(ax, 7.1, 2.5, 2.5, 1.3, 'LLM\nRetrieved Chunks\n+ Question', C_PRIMARY, 'white', 9)

    _arrow(ax, 9.6, 3.15, 10, 3.15)
    _box(ax, 9.7, 1.4, 2, 0.8, 'Answer with\nCitations', C_GREEN, 'white', 9)

    # Arrow from Qdrant store to similarity search
    ax.annotate('', xy=(5.8, 3.8), xytext=(8.1, 4.8),
                arrowprops=dict(arrowstyle='->', color=C_AMBER, lw=1.5, ls='--'))
    ax.text(7.6, 4.35, 'vector index', fontsize=7, color=C_AMBER, fontstyle='italic', rotation=-30)

    return _to_buf(fig)


def dia_search_flow():
    """Web search + grounded answer flow."""
    fig, ax = _fig(12, 4.5, (0, 12), (0, 4.5))
    ax.text(6, 4.2, 'Grounded Search Pipeline', ha='center', fontsize=13,
            fontweight='bold', color=C_DARK)

    steps = [
        (0.1, 'User\nQuery', C_TEAL),
        (2.3, 'SearXNG\nMeta-Search', C_ORANGE),
        (4.5, 'Top Results\nSnippets', C_AMBER),
        (6.7, 'LLM + Context\nGenerate Answer', C_PRIMARY),
        (9.2, 'Grounded\nResponse\n+ Citations', C_GREEN),
    ]
    bw, bh = 2, 1
    y = 1.5
    for x, label, color in steps:
        _box(ax, x, y, bw, bh, label, color, 'white', 9)
    for i in range(len(steps) - 1):
        _arrow(ax, steps[i][0] + bw, y + bh/2, steps[i+1][0], y + bh/2)

    # Cache note
    _box(ax, 4.5, 0.2, 2, 0.6, 'Result Cache\n(Redis)', '#e8e8e8', C_DARK, 7, 'normal')
    _arrow(ax, 5.5, 1.5, 5.5, 0.85, color=C_GRAY, lw=1, style='<->')

    return _to_buf(fig)


def dia_db_schema():
    """Database ER diagram."""
    fig, ax = _fig(11, 6, (0, 11), (0, 6))
    ax.text(5.5, 5.7, 'Database Schema', ha='center', fontsize=13, fontweight='bold', color=C_DARK)

    # Users table
    _box(ax, 0.3, 3.0, 3, 2.3, '', '#f0f4f8', C_DARK, 9)
    ax.text(1.8, 5.0, 'users', fontsize=10, fontweight='bold', color=C_PRIMARY, ha='center')
    fields_u = ['id (PK, UUID)', 'roll_number (UNIQUE)', 'name', 'department', 'role', 'hashed_password', 'is_active', 'created_at']
    for i, f in enumerate(fields_u):
        ax.text(0.5, 4.6 - i*0.2, f, fontsize=6.5, color=C_DARK, family='monospace')

    # api_keys table
    _box(ax, 4, 3.5, 3, 1.8, '', '#f0f4f8', C_DARK, 9)
    ax.text(5.5, 5.0, 'api_keys', fontsize=10, fontweight='bold', color=C_PRIMARY, ha='center')
    fields_k = ['id (PK, UUID)', 'user_id (FK)', 'key_hash', 'key_prefix', 'type', 'is_active', 'expires_at']
    for i, f in enumerate(fields_k):
        ax.text(4.2, 4.6 - i*0.2, f, fontsize=6.5, color=C_DARK, family='monospace')

    # request_logs table
    _box(ax, 7.7, 3.0, 3, 2.3, '', '#f0f4f8', C_DARK, 9)
    ax.text(9.2, 5.0, 'request_logs', fontsize=10, fontweight='bold', color=C_PRIMARY, ha='center')
    fields_r = ['id (PK, UUID)', 'user_id (FK)', 'api_key_id (FK)', 'model', 'endpoint', 'tokens_in', 'tokens_out', 'latency_ms', 'status_code']
    for i, f in enumerate(fields_r):
        ax.text(7.9, 4.6 - i*0.2, f, fontsize=6.5, color=C_DARK, family='monospace')

    # Relationships
    _arrow(ax, 3.3, 4.4, 4.0, 4.4, color=C_SEC, lw=2)
    ax.text(3.65, 4.55, '1:N', fontsize=7, color=C_SEC, ha='center', fontweight='bold')

    _arrow(ax, 3.3, 3.8, 7.7, 3.8, color=C_SEC, lw=2)
    ax.text(5.5, 3.95, '1:N', fontsize=7, color=C_SEC, ha='center', fontweight='bold')

    # quota_overrides
    _box(ax, 0.3, 0.5, 3, 1.3, '', '#f0f4f8', C_DARK, 9)
    ax.text(1.8, 1.55, 'quota_overrides', fontsize=9, fontweight='bold', color=C_PRIMARY, ha='center')
    q_fields = ['id (PK)', 'user_id (FK)', 'daily_tokens', 'requests_per_hr', 'set_by (FK)']
    for i, f in enumerate(q_fields):
        ax.text(0.5, 1.2 - i*0.2, f, fontsize=6.5, color=C_DARK, family='monospace')

    # rag_documents
    _box(ax, 7.7, 0.5, 3, 1.3, '', '#f0f4f8', C_DARK, 9)
    ax.text(9.2, 1.55, 'rag_documents', fontsize=9, fontweight='bold', color=C_PRIMARY, ha='center')
    d_fields = ['id (PK)', 'title', 'collection', 'chunk_count', 'uploaded_by (FK)']
    for i, f in enumerate(d_fields):
        ax.text(7.9, 1.2 - i*0.2, f, fontsize=6.5, color=C_DARK, family='monospace')

    _arrow(ax, 1.8, 3.0, 1.8, 1.85, color=C_SEC, lw=1.5)
    _arrow(ax, 3.3, 3.5, 7.7, 1.85, color=C_SEC, lw=1, style='->')

    return _to_buf(fig)


# ════════════════════════════════════════════════════════════════
#  DOCUMENT GENERATORS
# ════════════════════════════════════════════════════════════════

def gen_phase_0():
    """Phase 0 — Project Overview & Roadmap."""
    doc = new_doc('Project Overview & Build Roadmap', 'Complete Platform Blueprint — Phases 1 through 8')

    doc.add_heading('What is MAC?', level=1)
    body(doc, 'MAC (MBM AI Cloud) is a self-hosted AI inference platform that I am building for MBM Engineering College. '
         'The idea is straightforward — instead of paying for expensive cloud AI services like OpenAI or Google Vertex AI, '
         'we run the best open-source AI models on our own lab PCs and give every student access through a clean, '
         'well-documented REST API.')
    body(doc, 'Students connect from anywhere on the college LAN, authenticate with their roll number, receive a personal '
         'API key, and start querying AI models — for code generation, mathematical reasoning, image understanding, '
         'speech transcription, or general text tasks. No subscriptions, no per-token charges, no data leaving our network.')
    body(doc, 'The platform is designed to start on a single PC and scale horizontally. Once Phase 1 is proven on one '
         'machine, the exact same code runs across 5, 10, or 30 lab PCs simply by adding vLLM worker nodes — zero code changes required.')

    doc.add_heading('Why Build This?', level=2)
    body(doc, '1. Cost — Commercial API access for 500+ students is prohibitively expensive. Our hardware is already sitting idle after lab hours.')
    body(doc, '2. Privacy — Student queries and data never leave the college network. No third-party data processing.')
    body(doc, '3. Learning — Students interact with the actual infrastructure (Docker, GPUs, model serving) — this is hands-on AI engineering.')
    body(doc, '4. Availability — Works even when the internet is down. The entire stack runs offline after initial model download.')

    doc.add_heading('Build Roadmap', level=1)
    body(doc, 'The project is divided into 8 sequential phases. Each phase builds on the previous one and delivers a '
         'working increment. Here is the complete roadmap:')
    add_diagram(doc, dia_roadmap())

    styled_table(doc, ['Phase', 'Name', 'What It Delivers', 'Depends On'],
    [
        ['1', 'API Endpoints', 'Core REST API — auth, explore, query, usage', 'None'],
        ['2', 'LLM Models', 'Download, configure, and deploy 5 specialist models', 'Phase 1'],
        ['3', 'API–Model Integration', 'Wire models to endpoints via LiteLLM + vLLM', 'Phase 1, 2'],
        ['4', 'Usage Control', 'Rate limiting, token accounting, API key management', 'Phase 1'],
        ['5', 'Web Interface', 'Student dashboard, admin panel, playground', 'Phase 1, 4'],
        ['6', 'Guardrails', 'Input/output content filtering, safety checks', 'Phase 3'],
        ['7', 'Knowledgebase + RAG', 'Vector DB, document ingestion, retrieval chain', 'Phase 3'],
        ['8', 'Retrieval + Search', 'Self-hosted web search, grounded answers', 'Phase 3, 7'],
    ])

    doc.add_heading('System Architecture', level=1)
    body(doc, 'Below is the full system architecture showing how all components connect. Requests flow from the student '
         'client through Nginx (reverse proxy), to the FastAPI gateway (authentication, routing, rate limiting), then to '
         'LiteLLM (model routing and load balancing), and finally to vLLM workers that run the actual GPU inference.')
    add_diagram(doc, dia_architecture())

    doc.add_heading('Technology Stack', level=1)
    body(doc, 'I evaluated multiple options for each layer. Here is what I chose and why:')
    styled_table(doc, ['Layer', 'Technology', 'Why I Chose It'],
    [
        ['API Framework', 'FastAPI (Python 3.11+)', 'Async, auto-generates Swagger docs, type-safe, fastest Python framework'],
        ['Database', 'PostgreSQL 16', 'ACID-compliant, excellent JSON support, handles concurrent writes well'],
        ['Cache / Rate Limiter', 'Redis 7', 'In-memory, sub-millisecond lookups, built-in rate limiting primitives'],
        ['LLM Inference', 'vLLM', 'PagedAttention gives 2–4x throughput over naive serving; continuous batching'],
        ['Model Router', 'LiteLLM', 'Unified OpenAI-compatible proxy with load balancing and health checks'],
        ['Vector Database', 'Qdrant', 'Purpose-built for embeddings, fast similarity search, easy snapshots'],
        ['Web Search', 'SearXNG', 'Self-hosted meta-search engine, no API keys required, fully offline-capable'],
        ['Reverse Proxy', 'Nginx', 'TLS termination, request buffering, static file serving'],
        ['Containers', 'Docker + Compose', 'Reproducible deploys, one-command startup, easy scaling'],
        ['Frontend', 'React + Tailwind CSS', 'Component-driven, responsive, fast iteration'],
        ['Task Queue', 'Celery + Redis', 'Background jobs for document ingestion, model downloads'],
    ])

    doc.add_heading('Database Schema', level=1)
    body(doc, 'The relational schema covers users, API keys, request logs, quota overrides, and RAG documents. '
         'All primary keys use UUIDs for security (no sequential IDs to guess).')
    add_diagram(doc, dia_db_schema())

    doc.add_heading('Project Folder Structure', level=1)
    body(doc, 'The codebase follows a clean modular layout. Each router, model, and service has its own file.', size=10)
    code_block = (
        'mac/\n'
        '├── api/\n'
        '│   ├── main.py                 # FastAPI entry point\n'
        '│   ├── routers/                # One file per endpoint group\n'
        '│   │   ├── auth.py, explore.py, query.py, usage.py\n'
        '│   │   ├── keys.py, models.py, integration.py, quota.py\n'
        '│   │   └── guardrails.py, rag.py, search.py\n'
        '│   ├── models/                 # SQLAlchemy ORM models\n'
        '│   ├── schemas/                # Pydantic request/response schemas\n'
        '│   ├── core/                   # Auth, config, rate limiting, DB\n'
        '│   ├── services/               # Business logic layer\n'
        '│   └── requirements.txt\n'
        '├── litellm/config.yaml         # Model routing configuration\n'
        '├── nginx/nginx.conf            # Reverse proxy configuration\n'
        '├── frontend/                   # React dashboard\n'
        '├── scripts/                    # Utilities (seed users, download models)\n'
        '├── docker-compose.yml          # Full stack orchestration\n'
        '├── Dockerfile\n'
        '├── .env.example\n'
        '└── alembic/                    # Database migrations'
    )
    p = doc.add_paragraph()
    r = p.add_run(code_block)
    r.font.size = Pt(8)
    r.font.name = 'Courier New'

    doc.add_heading('Security Design', level=1)
    styled_table(doc, ['Concern', 'How It Is Handled'],
    [
        ['Password Storage', 'bcrypt hash with work factor 12; plaintext never stored'],
        ['JWT Signing', 'RS256 asymmetric keys; access tokens 24h, refresh tokens 30d'],
        ['API Key Storage', 'Only SHA-256 hash stored in database; raw key shown once at creation'],
        ['Transport', 'Nginx terminates TLS (HTTPS); internal comms over Docker network'],
        ['Input Validation', 'Pydantic schemas on every request; max payload size enforced by Nginx'],
        ['SQL Injection', 'SQLAlchemy ORM with parameterised queries'],
        ['Rate Limiting', 'Redis sliding-window per user + per IP'],
        ['CORS', 'Strict allowlist — only the MAC frontend origin'],
        ['Admin Routes', 'Role-based access control; JWT role claim checked in middleware'],
        ['Prompt Injection', 'Input guardrails detect and block override attempts'],
    ])

    doc.add_heading('OpenAI SDK Compatibility', level=1)
    body(doc, 'A core design decision: every /query endpoint returns OpenAI-compatible JSON. This means students '
         'can use the official OpenAI Python SDK — they just swap base_url. Any tutorial or blog post that uses '
         'the OpenAI API works instantly with MAC, no code rewriting needed.')
    code_ex = (
        'from openai import OpenAI\n\n'
        'client = OpenAI(\n'
        '    base_url="http://mac-server/api/v1",\n'
        '    api_key="mac_sk_live_your_key_here"\n'
        ')\n\n'
        'response = client.chat.completions.create(\n'
        '    model="auto",\n'
        '    messages=[{"role": "user", "content": "Explain binary search"}]\n'
        ')\n'
        'print(response.choices[0].message.content)'
    )
    p = doc.add_paragraph()
    r = p.add_run(code_ex)
    r.font.size = Pt(9)
    r.font.name = 'Courier New'

    save(doc, 'Phase-0-Project-Overview.docx')


def _json_block(doc, text):
    """Add a JSON code block to the document."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(8)
    r.font.name = 'Courier New'
    p.paragraph_format.space_after = Pt(8)

def gen_phase_1():
    """Phase 1 — API Endpoints — FULL PAYLOAD DETAILS."""
    doc = new_doc('Phase 1 — API Endpoints', 'explore . query . usage . auth — the foundation\nComplete Request/Response Payloads for Every Endpoint')

    doc.add_heading('Overview', level=1)
    body(doc, 'This is the foundation of the entire platform. In this phase, I am building four core endpoint groups '
         'that every other feature depends on. The API is designed to be OpenAI-compatible from the start, so students '
         'can use familiar SDKs and existing tutorials with zero modifications.')
    body(doc, 'Base URL: http://<server-ip>/api/v1 — the server address is configured via environment variable '
         '(MAC_HOST), never hardcoded. All endpoints below are relative to this base.')
    body(doc, 'Authentication: Endpoints under /query and /usage require either a JWT Bearer token (from /auth/login) '
         'or an API key in the Authorization header. Endpoints under /explore are public (no auth needed except /explore/usage-stats). '
         'All POST endpoints accept JSON request bodies with Content-Type: application/json.')

    # ── AUTH ────────────────────────────────────────────────
    doc.add_heading('Authentication Flow', level=1)
    body(doc, 'Students log in with their roll number and password. The server validates credentials against the '
         'bcrypt-hashed password in PostgreSQL, generates a JWT access token (RS256 signed, 24-hour expiry) and a refresh token '
         '(30-day expiry), and returns both along with the student profile and API key.')
    body(doc, 'Why JWT? They are stateless — the server does not need to hit the database on every request. Any server '
         'with the public key can verify the token. This is important for horizontal scaling.')
    add_diagram(doc, dia_auth_flow())

    doc.add_heading('1.1 Authentication Endpoints — /auth', level=1)
    styled_table(doc, ['Method', 'Endpoint', 'Auth Required', 'Description'],
    [
        ['POST', '/auth/login', 'No', 'Roll number + password -> JWT access token + refresh token'],
        ['POST', '/auth/logout', 'Yes (JWT)', 'Invalidate current session / revoke refresh token'],
        ['POST', '/auth/refresh', 'No (uses refresh token)', 'Exchange refresh token for new access token'],
        ['GET', '/auth/me', 'Yes (JWT)', 'Get current user profile, role, department, API key'],
        ['POST', '/auth/change-password', 'Yes (JWT)', 'Change password (requires current password)'],
    ])

    # --- POST /auth/login ---
    doc.add_heading('POST /auth/login', level=2)
    body(doc, 'Authenticate a student with roll number and password. Returns JWT tokens and user profile.', italic=True)
    body(doc, 'Request Headers:', bold=True, size=10)
    styled_table(doc, ['Header', 'Value'],
    [['Content-Type', 'application/json']])
    body(doc, 'Request Body:', bold=True, size=10)
    styled_table(doc, ['Field', 'Type', 'Required', 'Constraints', 'Description'],
    [
        ['roll_number', 'string', 'Yes', 'Pattern: [0-9]{2}[A-Z]{2,4}[0-9]{3}', 'Student roll number e.g. 21CS045'],
        ['password', 'string', 'Yes', 'Min 8 chars', 'Account password'],
    ])
    body(doc, 'Example Request:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "roll_number": "21CS045",\n'
        '  "password": "mySecurePass123"\n'
        '}')
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    styled_table(doc, ['Field', 'Type', 'Description'],
    [
        ['access_token', 'string', 'JWT token, RS256 signed, 24-hour expiry'],
        ['refresh_token', 'string', 'Opaque token, 30-day expiry, stored hashed in DB'],
        ['token_type', 'string', 'Always "bearer"'],
        ['expires_in', 'integer', 'Seconds until access_token expires (86400 = 24h)'],
        ['user', 'object', 'User profile object (see below)'],
        ['user.roll_number', 'string', 'Student roll number'],
        ['user.name', 'string', 'Full name'],
        ['user.department', 'string', 'Department code e.g. CSE, ECE, ME'],
        ['user.role', 'string', 'One of: student, faculty, admin'],
        ['user.api_key', 'string', 'Personal API key (mac_sk_live_xxx...)'],
    ], col_widths=[2.0, 0.8, 3.8])
    body(doc, 'Example Response:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "access_token": "eyJhbGciOiJSUzI1NiIsInR5cCI6IkpXVCJ9...",\n'
        '  "refresh_token": "dGhpcyBpcyBhIHJlZnJlc2ggdG9rZW4...",\n'
        '  "token_type": "bearer",\n'
        '  "expires_in": 86400,\n'
        '  "user": {\n'
        '    "roll_number": "21CS045",\n'
        '    "name": "Aryan Sharma",\n'
        '    "department": "CSE",\n'
        '    "role": "student",\n'
        '    "api_key": "mac_sk_live_a1b2c3d4e5f6..."\n'
        '  }\n'
        '}')
    body(doc, 'Error Responses:', bold=True, size=10)
    styled_table(doc, ['Status', 'Code', 'When'],
    [
        ['401', 'authentication_failed', 'Wrong roll number or password'],
        ['422', 'validation_error', 'Missing required fields or invalid format'],
        ['423', 'account_locked', 'Account locked after 5 failed attempts (15 min cooldown)'],
    ])

    # --- POST /auth/logout ---
    doc.add_heading('POST /auth/logout', level=2)
    body(doc, 'Invalidate the current session. The refresh token is revoked in the database.', italic=True)
    body(doc, 'Request Headers:', bold=True, size=10)
    styled_table(doc, ['Header', 'Value'],
    [['Authorization', 'Bearer <access_token>']])
    body(doc, 'Request Body: None (empty)', size=10)
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc, '{\n  "message": "Successfully logged out"\n}')

    # --- POST /auth/refresh ---
    doc.add_heading('POST /auth/refresh', level=2)
    body(doc, 'Exchange a valid refresh token for a new access token without re-entering password.', italic=True)
    body(doc, 'Request Body:', bold=True, size=10)
    styled_table(doc, ['Field', 'Type', 'Required', 'Description'],
    [['refresh_token', 'string', 'Yes', 'The refresh token received from /auth/login']])
    body(doc, 'Example Request:', bold=True, size=10)
    _json_block(doc, '{\n  "refresh_token": "dGhpcyBpcyBhIHJlZnJlc2ggdG9rZW4..."\n}')
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "access_token": "eyJhbGciOiJSUzI1NiJ9...(new token)...",\n'
        '  "token_type": "bearer",\n'
        '  "expires_in": 86400\n'
        '}')
    body(doc, 'Error: 401 if refresh token is expired or revoked.', size=10)

    # --- GET /auth/me ---
    doc.add_heading('GET /auth/me', level=2)
    body(doc, 'Return the current authenticated user\'s full profile.', italic=True)
    body(doc, 'Request Headers:', bold=True, size=10)
    styled_table(doc, ['Header', 'Value'],
    [['Authorization', 'Bearer <access_token>']])
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "roll_number": "21CS045",\n'
        '  "name": "Aryan Sharma",\n'
        '  "department": "CSE",\n'
        '  "role": "student",\n'
        '  "is_active": true,\n'
        '  "api_key": "mac_sk_live_a1b2...****",\n'
        '  "api_key_type": "refresh",\n'
        '  "quota": {\n'
        '    "daily_tokens": 50000,\n'
        '    "tokens_used_today": 12450,\n'
        '    "requests_per_hour": 100,\n'
        '    "requests_this_hour": 8\n'
        '  },\n'
        '  "created_at": "2026-01-15T10:00:00Z"\n'
        '}')

    # --- POST /auth/change-password ---
    doc.add_heading('POST /auth/change-password', level=2)
    body(doc, 'Change the current user\'s password. Requires the old password for verification.', italic=True)
    body(doc, 'Request Body:', bold=True, size=10)
    styled_table(doc, ['Field', 'Type', 'Required', 'Description'],
    [
        ['old_password', 'string', 'Yes', 'Current password for verification'],
        ['new_password', 'string', 'Yes', 'New password (min 8 chars, must contain letter + digit)'],
    ])
    body(doc, 'Example Request:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "old_password": "mySecurePass123",\n'
        '  "new_password": "evenMoreSecure456"\n'
        '}')
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc, '{\n  "message": "Password changed successfully"\n}')
    body(doc, 'Error: 401 if old_password is incorrect. 422 if new_password is too weak.', size=10)

    # ── EXPLORE ────────────────────────────────────────────
    doc.add_heading('1.2 Explore Endpoints — /explore', level=1)
    body(doc, 'Read-only discovery endpoints. These let students see what models are available, what capabilities '
         'the platform offers, and whether the system is healthy — all before writing any code. '
         'No authentication required except for /explore/usage-stats (admin-only).')
    styled_table(doc, ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ['GET', '/explore/models', 'No', 'List all deployed models with capabilities, context length, status'],
        ['GET', '/explore/models/{model_id}', 'No', 'Detailed model info — parameters, benchmarks, example prompts'],
        ['GET', '/explore/models/search', 'No', 'Search by capability tag: ?tag=vision, ?tag=code, ?tag=math'],
        ['GET', '/explore/endpoints', 'No', 'List every API endpoint with method, path, auth requirement'],
        ['GET', '/explore/health', 'No', 'Platform health — node status, GPU temps, queue depth'],
        ['GET', '/explore/usage-stats', 'Admin only', 'Aggregated platform analytics — tokens/day, active users'],
    ])

    # --- GET /explore/models ---
    doc.add_heading('GET /explore/models', level=2)
    body(doc, 'Returns a list of all models deployed on the platform with their current status.', italic=True)
    body(doc, 'Query Parameters (all optional):', bold=True, size=10)
    styled_table(doc, ['Param', 'Type', 'Default', 'Description'],
    [
        ['status', 'string', 'all', 'Filter by status: loaded, offline, downloading'],
        ['capability', 'string', 'all', 'Filter by capability: code, chat, vision, speech, math'],
        ['page', 'int', '1', 'Page number for pagination'],
        ['per_page', 'int', '20', 'Items per page (max 100)'],
    ])
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "models": [\n'
        '    {\n'
        '      "id": "qwen2.5-coder-7b",\n'
        '      "name": "Qwen2.5-Coder 7B",\n'
        '      "specialty": "Code generation, debugging, explanation",\n'
        '      "parameters": "7B",\n'
        '      "context_length": 32768,\n'
        '      "quantisation": "GPTQ-Int4",\n'
        '      "vram_mb": 5120,\n'
        '      "status": "loaded",\n'
        '      "capabilities": ["code", "chat", "completion"],\n'
        '      "loaded_at": "2026-04-07T08:30:00Z"\n'
        '    },\n'
        '    {\n'
        '      "id": "deepseek-r1-8b",\n'
        '      "name": "DeepSeek-R1 8B",\n'
        '      "specialty": "Maths, reasoning, step-by-step logic",\n'
        '      "parameters": "8B",\n'
        '      "context_length": 65536,\n'
        '      "quantisation": "AWQ-Int4",\n'
        '      "vram_mb": 6144,\n'
        '      "status": "offline",\n'
        '      "capabilities": ["reasoning", "math", "chat"],\n'
        '      "loaded_at": null\n'
        '    }\n'
        '  ],\n'
        '  "total": 5,\n'
        '  "page": 1,\n'
        '  "per_page": 20\n'
        '}')

    # --- GET /explore/models/{model_id} ---
    doc.add_heading('GET /explore/models/{model_id}', level=2)
    body(doc, 'Return detailed information about a single model including benchmarks and example usage.', italic=True)
    body(doc, 'Path Parameters:', bold=True, size=10)
    styled_table(doc, ['Param', 'Type', 'Description'],
    [['model_id', 'string', 'Model identifier e.g. qwen2.5-coder-7b']])
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "id": "qwen2.5-coder-7b",\n'
        '  "name": "Qwen2.5-Coder 7B",\n'
        '  "specialty": "Code generation, debugging, explanation",\n'
        '  "parameters": "7B",\n'
        '  "context_length": 32768,\n'
        '  "quantisation": "GPTQ-Int4",\n'
        '  "vram_mb": 5120,\n'
        '  "status": "loaded",\n'
        '  "capabilities": ["code", "chat", "completion"],\n'
        '  "benchmarks": {\n'
        '    "humaneval": 85.4,\n'
        '    "mbpp": 78.2\n'
        '  },\n'
        '  "example_prompt": "Write a Python function to reverse a linked list",\n'
        '  "supported_languages": ["python", "javascript", "java", "c++", "rust"],\n'
        '  "loaded_at": "2026-04-07T08:30:00Z",\n'
        '  "total_requests_served": 14520\n'
        '}')
    body(doc, 'Error: 404 if model_id does not exist.', size=10)

    # --- GET /explore/models/search ---
    doc.add_heading('GET /explore/models/search', level=2)
    body(doc, 'Search models by capability tag.', italic=True)
    body(doc, 'Query Parameters:', bold=True, size=10)
    styled_table(doc, ['Param', 'Type', 'Required', 'Description'],
    [['tag', 'string', 'Yes', 'Capability to search for: vision, code, math, speech, chat, reasoning']])
    body(doc, 'Example: GET /explore/models/search?tag=code', size=10)
    body(doc, 'Response: Same structure as GET /explore/models, filtered by matching capability.', size=10)

    # --- GET /explore/endpoints ---
    doc.add_heading('GET /explore/endpoints', level=2)
    body(doc, 'List every available API endpoint. Useful for client auto-discovery.', italic=True)
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "endpoints": [\n'
        '    {\n'
        '      "method": "POST",\n'
        '      "path": "/api/v1/auth/login",\n'
        '      "auth_required": false,\n'
        '      "description": "Authenticate with roll number and password",\n'
        '      "request_content_type": "application/json"\n'
        '    },\n'
        '    {\n'
        '      "method": "POST",\n'
        '      "path": "/api/v1/query/chat",\n'
        '      "auth_required": true,\n'
        '      "description": "Chat completion - multi-turn conversation",\n'
        '      "request_content_type": "application/json"\n'
        '    }\n'
        '  ],\n'
        '  "total": 35\n'
        '}')

    # --- GET /explore/health ---
    doc.add_heading('GET /explore/health', level=2)
    body(doc, 'Platform health check — shows node status, GPU temperatures, and inference queue depth.', italic=True)
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "status": "healthy",\n'
        '  "uptime_seconds": 345600,\n'
        '  "version": "1.0.0",\n'
        '  "nodes": [\n'
        '    {\n'
        '      "id": "node-01",\n'
        '      "ip": "192.168.1.101",\n'
        '      "gpu": "NVIDIA RTX 3060 12GB",\n'
        '      "gpu_temp_c": 62,\n'
        '      "vram_used_gb": 9.2,\n'
        '      "vram_total_gb": 12.0,\n'
        '      "models_loaded": ["qwen2.5-coder-7b"],\n'
        '      "requests_in_flight": 3,\n'
        '      "status": "active"\n'
        '    }\n'
        '  ],\n'
        '  "queue_depth": 7,\n'
        '  "models_loaded": 3,\n'
        '  "models_total": 5\n'
        '}')

    # --- GET /explore/usage-stats ---
    doc.add_heading('GET /explore/usage-stats (admin-only)', level=2)
    body(doc, 'Aggregated platform usage statistics. Requires admin role.', italic=True)
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "today": {\n'
        '    "total_tokens": 2450000,\n'
        '    "total_requests": 4870,\n'
        '    "unique_users": 127,\n'
        '    "avg_latency_ms": 340\n'
        '  },\n'
        '  "this_week": {\n'
        '    "total_tokens": 12800000,\n'
        '    "total_requests": 28450,\n'
        '    "unique_users": 312\n'
        '  },\n'
        '  "top_models": [\n'
        '    {"model": "qwen2.5-14b", "requests": 2100, "tokens": 1200000},\n'
        '    {"model": "qwen2.5-coder-7b", "requests": 1800, "tokens": 850000}\n'
        '  ],\n'
        '  "peak_hour": "14:00-15:00"\n'
        '}')

    # ── QUERY ──────────────────────────────────────────────
    doc.add_heading('1.3 Query Endpoints — /query', level=1)
    body(doc, 'This is the core inference API. All endpoints require authentication via either: '
         '(1) Authorization: Bearer <access_token> (JWT from login) or '
         '(2) Authorization: Bearer <api_key> (mac_sk_live_xxx). '
         'All responses follow the OpenAI Chat Completions format — any OpenAI SDK works by swapping base_url.')
    styled_table(doc, ['Method', 'Endpoint', 'Content-Type', 'Description'],
    [
        ['POST', '/query/chat', 'application/json', 'Chat completion — multi-turn (text/code/math)'],
        ['POST', '/query/completions', 'application/json', 'Raw text completion (OpenAI-compatible)'],
        ['POST', '/query/vision', 'multipart/form-data', 'Image + text -> answer'],
        ['POST', '/query/speech-to-text', 'multipart/form-data', 'Audio file -> transcribed text'],
        ['POST', '/query/text-to-speech', 'application/json', 'Text -> audio file download'],
        ['POST', '/query/embeddings', 'application/json', 'Text -> vector embedding'],
        ['POST', '/query/rerank', 'application/json', 'Re-rank passages by relevance'],
    ])

    # --- POST /query/chat ---
    doc.add_heading('POST /query/chat', level=2)
    body(doc, 'The primary inference endpoint. Supports multi-turn conversation, streaming, and smart model routing.', italic=True)
    body(doc, 'Request Headers:', bold=True, size=10)
    styled_table(doc, ['Header', 'Value'],
    [
        ['Content-Type', 'application/json'],
        ['Authorization', 'Bearer <access_token> or Bearer <api_key>'],
    ])
    body(doc, 'Request Body:', bold=True, size=10)
    styled_table(doc, ['Field', 'Type', 'Required', 'Default', 'Constraints', 'Description'],
    [
        ['model', 'string', 'Yes', '-', 'Valid model ID or "auto"', 'Model to use. "auto" = smart routing based on content'],
        ['messages', 'array', 'Yes', '-', 'Min 1 message object', 'Array of {role, content} objects'],
        ['messages[].role', 'string', 'Yes', '-', 'system|user|assistant', 'Role of this message in the conversation'],
        ['messages[].content', 'string', 'Yes', '-', 'Max 32000 chars', 'The message text content'],
        ['temperature', 'float', 'No', '0.7', '0.0 - 2.0', 'Higher = more creative, lower = more focused'],
        ['max_tokens', 'integer', 'No', '2048', '1 - 4096 (student), 8192 (faculty)', 'Max tokens to generate'],
        ['stream', 'boolean', 'No', 'false', '-', 'If true, response is Server-Sent Events stream'],
        ['top_p', 'float', 'No', '1.0', '0.0 - 1.0', 'Nucleus sampling: only consider top_p probability mass'],
        ['frequency_penalty', 'float', 'No', '0.0', '-2.0 - 2.0', 'Penalise repeated tokens'],
        ['presence_penalty', 'float', 'No', '0.0', '-2.0 - 2.0', 'Penalise tokens already in context'],
        ['stop', 'array|string', 'No', 'null', 'Max 4 sequences', 'Stop generating when these sequences appear'],
        ['context_id', 'string', 'No', 'null', 'UUID format', 'Resume a previous conversation server-side'],
    ], col_widths=[1.5, 0.7, 0.6, 0.5, 1.5, 1.8])
    body(doc, 'Example Request:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "model": "auto",\n'
        '  "messages": [\n'
        '    {"role": "system", "content": "You are a helpful coding assistant."},\n'
        '    {"role": "user", "content": "Write a Python function to reverse a linked list"}\n'
        '  ],\n'
        '  "temperature": 0.7,\n'
        '  "max_tokens": 2048,\n'
        '  "stream": false\n'
        '}')
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    styled_table(doc, ['Field', 'Type', 'Description'],
    [
        ['id', 'string', 'Unique response ID (mac-chat-xxxx)'],
        ['object', 'string', 'Always "chat.completion"'],
        ['created', 'integer', 'Unix timestamp of generation'],
        ['model', 'string', 'Actual model used (useful when model="auto")'],
        ['choices', 'array', 'Array of completion choices (usually 1)'],
        ['choices[].index', 'integer', 'Choice index'],
        ['choices[].message.role', 'string', 'Always "assistant"'],
        ['choices[].message.content', 'string', 'The generated text response'],
        ['choices[].finish_reason', 'string', '"stop" (natural end) or "length" (hit max_tokens)'],
        ['usage', 'object', 'Token usage stats'],
        ['usage.prompt_tokens', 'integer', 'Tokens in the input'],
        ['usage.completion_tokens', 'integer', 'Tokens generated'],
        ['usage.total_tokens', 'integer', 'Sum of prompt + completion'],
        ['context_id', 'string', 'Conversation ID for follow-up calls'],
    ], col_widths=[2.2, 0.8, 3.6])
    body(doc, 'Example Response:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "id": "mac-chat-a1b2c3d4",\n'
        '  "object": "chat.completion",\n'
        '  "created": 1743984000,\n'
        '  "model": "qwen2.5-coder-7b",\n'
        '  "choices": [\n'
        '    {\n'
        '      "index": 0,\n'
        '      "message": {\n'
        '        "role": "assistant",\n'
        '        "content": "Here is a Python function to reverse a linked list...\\n\\n```python\\nclass ListNode:\\n    def __init__(self, val=0):\\n        self.val = val\\n        self.next = None\\n\\ndef reverse_list(head):\\n    prev = None\\n    current = head\\n    while current:\\n        next_node = current.next\\n        current.next = prev\\n        prev = current\\n        current = next_node\\n    return prev\\n```"\n'
        '      },\n'
        '      "finish_reason": "stop"\n'
        '    }\n'
        '  ],\n'
        '  "usage": {\n'
        '    "prompt_tokens": 42,\n'
        '    "completion_tokens": 187,\n'
        '    "total_tokens": 229\n'
        '  },\n'
        '  "context_id": "ctx-e5f6g7h8"\n'
        '}')

    body(doc, 'Streaming Response (when stream=true):', bold=True, size=10)
    body(doc, 'When stream is true, the response is delivered as Server-Sent Events (SSE). Each event contains a delta with '
         'a partial token. The client reads chunks as they arrive for real-time display:', size=10)
    _json_block(doc,
        'data: {"id":"mac-chat-a1b2","choices":[{"delta":{"content":"Here"},"index":0}]}\n'
        'data: {"id":"mac-chat-a1b2","choices":[{"delta":{"content":" is"},"index":0}]}\n'
        'data: {"id":"mac-chat-a1b2","choices":[{"delta":{"content":" a"},"index":0}]}\n'
        '...\n'
        'data: [DONE]')

    # --- POST /query/completions ---
    doc.add_heading('POST /query/completions', level=2)
    body(doc, 'Raw text completion (non-chat format). Compatible with OpenAI Completions API.', italic=True)
    body(doc, 'Request Body:', bold=True, size=10)
    styled_table(doc, ['Field', 'Type', 'Required', 'Default', 'Description'],
    [
        ['model', 'string', 'Yes', '-', 'Model ID or "auto"'],
        ['prompt', 'string', 'Yes', '-', 'The text prompt to complete'],
        ['max_tokens', 'integer', 'No', '256', 'Max tokens to generate'],
        ['temperature', 'float', 'No', '0.7', 'Sampling temperature'],
        ['stop', 'array|string', 'No', 'null', 'Stop sequences'],
        ['echo', 'boolean', 'No', 'false', 'Include prompt in the response'],
    ])
    body(doc, 'Example Request:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "model": "qwen2.5-14b",\n'
        '  "prompt": "The capital of France is",\n'
        '  "max_tokens": 50,\n'
        '  "temperature": 0.3\n'
        '}')
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "id": "mac-comp-x1y2z3",\n'
        '  "object": "text_completion",\n'
        '  "created": 1743984100,\n'
        '  "model": "qwen2.5-14b",\n'
        '  "choices": [\n'
        '    {\n'
        '      "text": " Paris. It is known as the City of Light...",\n'
        '      "index": 0,\n'
        '      "finish_reason": "stop"\n'
        '    }\n'
        '  ],\n'
        '  "usage": {"prompt_tokens": 7, "completion_tokens": 15, "total_tokens": 22}\n'
        '}')

    # --- POST /query/vision ---
    doc.add_heading('POST /query/vision', level=2)
    body(doc, 'Send an image along with a text question. The vision model analyses the image and responds.', italic=True)
    body(doc, 'Request: multipart/form-data', bold=True, size=10)
    styled_table(doc, ['Field', 'Type', 'Required', 'Description'],
    [
        ['image', 'file', 'Yes', 'Image file (JPEG, PNG, WebP). Max 10 MB.'],
        ['question', 'string', 'Yes', 'The question about the image'],
        ['model', 'string', 'No', 'Default: llava-1.6-7b'],
        ['max_tokens', 'integer', 'No', 'Default: 1024'],
    ])
    body(doc, 'Example (curl):', bold=True, size=10)
    _json_block(doc,
        'curl -X POST http://mac-server/api/v1/query/vision \\\n'
        '  -H "Authorization: Bearer mac_sk_live_xxx" \\\n'
        '  -F "image=@circuit_diagram.png" \\\n'
        '  -F "question=What components are shown in this circuit?"')
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "id": "mac-vis-m1n2o3",\n'
        '  "model": "llava-1.6-7b",\n'
        '  "answer": "The circuit diagram shows a series connection with...",\n'
        '  "usage": {"prompt_tokens": 580, "completion_tokens": 120, "total_tokens": 700}\n'
        '}')

    # --- POST /query/speech-to-text ---
    doc.add_heading('POST /query/speech-to-text', level=2)
    body(doc, 'Transcribe an audio file to text using Whisper.', italic=True)
    body(doc, 'Request: multipart/form-data', bold=True, size=10)
    styled_table(doc, ['Field', 'Type', 'Required', 'Description'],
    [
        ['audio', 'file', 'Yes', 'Audio file (WAV, MP3, FLAC, M4A). Max 25 MB.'],
        ['language', 'string', 'No', 'ISO language code e.g. "en", "hi". Auto-detected if omitted.'],
        ['response_format', 'string', 'No', '"json" (default), "text", "srt", "vtt"'],
    ])
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "id": "mac-stt-p1q2r3",\n'
        '  "model": "whisper-large-v3",\n'
        '  "text": "Good morning class, today we will study binary search trees...",\n'
        '  "language": "en",\n'
        '  "duration_seconds": 45.2,\n'
        '  "segments": [\n'
        '    {"start": 0.0, "end": 3.2, "text": "Good morning class,"},\n'
        '    {"start": 3.2, "end": 7.1, "text": "today we will study binary search trees..."}\n'
        '  ]\n'
        '}')

    # --- POST /query/text-to-speech ---
    doc.add_heading('POST /query/text-to-speech', level=2)
    body(doc, 'Convert text to spoken audio. Returns a downloadable audio file.', italic=True)
    body(doc, 'Request Body:', bold=True, size=10)
    styled_table(doc, ['Field', 'Type', 'Required', 'Description'],
    [
        ['text', 'string', 'Yes', 'The text to convert to speech (max 4096 chars)'],
        ['voice', 'string', 'No', 'Voice ID (default: "default"). Options vary by TTS model.'],
        ['speed', 'float', 'No', 'Speed multiplier 0.5-2.0 (default: 1.0)'],
        ['response_format', 'string', 'No', '"mp3" (default), "wav", "opus"'],
    ])
    body(doc, 'Example Request:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "text": "Binary search works by repeatedly dividing the search interval in half.",\n'
        '  "voice": "default",\n'
        '  "speed": 1.0,\n'
        '  "response_format": "mp3"\n'
        '}')
    body(doc, 'Response: Binary audio file with Content-Type: audio/mpeg (or audio/wav, audio/opus).', size=10)

    # --- POST /query/embeddings ---
    doc.add_heading('POST /query/embeddings', level=2)
    body(doc, 'Generate vector embeddings for text. Used for RAG, semantic search, and similarity comparisons.', italic=True)
    body(doc, 'Request Body:', bold=True, size=10)
    styled_table(doc, ['Field', 'Type', 'Required', 'Description'],
    [
        ['input', 'string|array', 'Yes', 'Text or array of texts to embed (max 8192 tokens per text)'],
        ['model', 'string', 'No', 'Embedding model (default uses platform default)'],
    ])
    body(doc, 'Example Request:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "input": ["Binary search tree", "Hash table lookup"],\n'
        '  "model": "default"\n'
        '}')
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "object": "list",\n'
        '  "data": [\n'
        '    {\n'
        '      "object": "embedding",\n'
        '      "index": 0,\n'
        '      "embedding": [0.0023, -0.0145, 0.0367, ...768 floats...]\n'
        '    },\n'
        '    {\n'
        '      "object": "embedding",\n'
        '      "index": 1,\n'
        '      "embedding": [0.0112, -0.0089, 0.0234, ...768 floats...]\n'
        '    }\n'
        '  ],\n'
        '  "model": "bge-base-en-v1.5",\n'
        '  "usage": {"prompt_tokens": 8, "total_tokens": 8}\n'
        '}')

    # --- POST /query/rerank ---
    doc.add_heading('POST /query/rerank', level=2)
    body(doc, 'Re-rank a list of text passages by relevance to a query. Useful for improving RAG retrieval quality.', italic=True)
    body(doc, 'Request Body:', bold=True, size=10)
    styled_table(doc, ['Field', 'Type', 'Required', 'Description'],
    [
        ['query', 'string', 'Yes', 'The reference query'],
        ['documents', 'array', 'Yes', 'Array of text strings to rank'],
        ['top_k', 'integer', 'No', 'Return only top-k results (default: all)'],
    ])
    body(doc, 'Example Request:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "query": "How does quicksort work?",\n'
        '  "documents": [\n'
        '    "Quicksort is a divide-and-conquer algorithm...",\n'
        '    "Merge sort splits the array into halves...",\n'
        '    "Quicksort picks a pivot element and partitions..."\n'
        '  ],\n'
        '  "top_k": 2\n'
        '}')
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "results": [\n'
        '    {"index": 2, "document": "Quicksort picks a pivot...", "relevance_score": 0.94},\n'
        '    {"index": 0, "document": "Quicksort is a divide-and-conquer...", "relevance_score": 0.91}\n'
        '  ]\n'
        '}')

    # ── USAGE ──────────────────────────────────────────────
    doc.add_heading('1.4 Usage Endpoints — /usage', level=1)
    body(doc, 'Every API call is logged with token counts, the model used, latency, and a timestamp. These endpoints '
         'let students track their own consumption and let admins monitor the whole platform. '
         'All endpoints require authentication.')
    styled_table(doc, ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ['GET', '/usage/me', 'Any user', 'My token usage - today, this week, this month, by model'],
        ['GET', '/usage/me/history', 'Any user', 'Full request history with pagination'],
        ['GET', '/usage/me/quota', 'Any user', 'My current quota limits and remaining balance'],
        ['GET', '/usage/admin/all', 'Admin only', 'All users usage summary'],
        ['GET', '/usage/admin/user/{roll_no}', 'Admin only', 'Specific student\'s full usage details'],
        ['GET', '/usage/admin/models', 'Admin only', 'Per-model usage stats across the platform'],
    ])

    # --- GET /usage/me ---
    doc.add_heading('GET /usage/me', level=2)
    body(doc, 'Return the current user\'s token usage broken down by time period and model.', italic=True)
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "roll_number": "21CS045",\n'
        '  "usage": {\n'
        '    "today": {\n'
        '      "total_tokens": 12450,\n'
        '      "prompt_tokens": 4200,\n'
        '      "completion_tokens": 8250,\n'
        '      "requests": 23,\n'
        '      "by_model": {\n'
        '        "qwen2.5-coder-7b": {"tokens": 8200, "requests": 15},\n'
        '        "qwen2.5-14b": {"tokens": 4250, "requests": 8}\n'
        '      }\n'
        '    },\n'
        '    "this_week": {"total_tokens": 67800, "requests": 142},\n'
        '    "this_month": {"total_tokens": 234500, "requests": 487}\n'
        '  },\n'
        '  "quota": {\n'
        '    "daily_limit": 50000,\n'
        '    "remaining_today": 37550,\n'
        '    "resets_at": "2026-04-08T00:00:00Z"\n'
        '  }\n'
        '}')

    # --- GET /usage/me/history ---
    doc.add_heading('GET /usage/me/history', level=2)
    body(doc, 'Paginated history of all requests made by the current user.', italic=True)
    body(doc, 'Query Parameters:', bold=True, size=10)
    styled_table(doc, ['Param', 'Type', 'Default', 'Description'],
    [
        ['page', 'int', '1', 'Page number'],
        ['per_page', 'int', '50', 'Items per page (max 100)'],
        ['model', 'string', 'all', 'Filter by model ID'],
        ['date_from', 'string', '-', 'ISO date filter start e.g. 2026-04-01'],
        ['date_to', 'string', '-', 'ISO date filter end e.g. 2026-04-07'],
    ])
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "requests": [\n'
        '    {\n'
        '      "id": "req-a1b2c3d4",\n'
        '      "model": "qwen2.5-coder-7b",\n'
        '      "endpoint": "/query/chat",\n'
        '      "tokens_in": 42,\n'
        '      "tokens_out": 187,\n'
        '      "latency_ms": 2340,\n'
        '      "status_code": 200,\n'
        '      "created_at": "2026-04-07T14:23:00Z"\n'
        '    }\n'
        '  ],\n'
        '  "total": 487,\n'
        '  "page": 1,\n'
        '  "per_page": 50\n'
        '}')

    # --- GET /usage/me/quota ---
    doc.add_heading('GET /usage/me/quota', level=2)
    body(doc, 'Show the current user\'s quota limits and how much has been consumed.', italic=True)
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "role": "student",\n'
        '  "limits": {\n'
        '    "daily_tokens": 50000,\n'
        '    "requests_per_hour": 100,\n'
        '    "max_tokens_per_request": 4096\n'
        '  },\n'
        '  "current": {\n'
        '    "tokens_used_today": 12450,\n'
        '    "requests_this_hour": 8,\n'
        '    "remaining_tokens": 37550,\n'
        '    "remaining_requests": 92\n'
        '  },\n'
        '  "resets": {\n'
        '    "daily_reset": "2026-04-08T00:00:00Z",\n'
        '    "hourly_reset": "2026-04-07T15:00:00Z"\n'
        '  },\n'
        '  "has_override": false\n'
        '}')

    # --- Admin endpoints ---
    doc.add_heading('GET /usage/admin/all (admin-only)', level=2)
    body(doc, 'Summary of all users\' usage. Paginated, sortable.', italic=True)
    body(doc, 'Query Parameters:', bold=True, size=10)
    styled_table(doc, ['Param', 'Type', 'Default', 'Description'],
    [
        ['page', 'int', '1', 'Page number'],
        ['per_page', 'int', '50', 'Items per page'],
        ['sort_by', 'string', 'tokens_today', 'Sort field: tokens_today, requests_today, name'],
        ['order', 'string', 'desc', 'Sort order: asc, desc'],
        ['department', 'string', 'all', 'Filter by department'],
    ])
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "users": [\n'
        '    {\n'
        '      "roll_number": "21CS045",\n'
        '      "name": "Aryan Sharma",\n'
        '      "department": "CSE",\n'
        '      "tokens_today": 12450,\n'
        '      "requests_today": 23,\n'
        '      "quota_used_pct": 24.9,\n'
        '      "last_active": "2026-04-07T14:23:00Z"\n'
        '    }\n'
        '  ],\n'
        '  "total_users": 312,\n'
        '  "page": 1\n'
        '}')

    doc.add_heading('GET /usage/admin/user/{roll_number} (admin-only)', level=2)
    body(doc, 'Full usage details for a specific student.', italic=True)
    body(doc, 'Path Parameter: roll_number (string) — e.g. 21CS045', size=10)
    body(doc, 'Response: Same structure as GET /usage/me but for the specified student, plus their account info.', size=10)

    doc.add_heading('GET /usage/admin/models (admin-only)', level=2)
    body(doc, 'Usage statistics broken down per model.', italic=True)
    body(doc, 'Response — 200 OK:', bold=True, size=10)
    _json_block(doc,
        '{\n'
        '  "models": [\n'
        '    {\n'
        '      "model_id": "qwen2.5-coder-7b",\n'
        '      "requests_today": 1800,\n'
        '      "tokens_today": 850000,\n'
        '      "avg_latency_ms": 280,\n'
        '      "unique_users_today": 87,\n'
        '      "peak_hour": "14:00-15:00",\n'
        '      "error_rate_pct": 0.3\n'
        '    }\n'
        '  ]\n'
        '}')

    # ── COMMON ─────────────────────────────────────────────
    doc.add_heading('Common Response Headers', level=1)
    body(doc, 'Every authenticated API response includes these headers:', size=10)
    styled_table(doc, ['Header', 'Example', 'Description'],
    [
        ['X-Request-ID', 'mac-req-a1b2c3', 'Unique ID for this request (useful for debugging)'],
        ['X-RateLimit-Limit', '100', 'Max requests allowed per hour'],
        ['X-RateLimit-Remaining', '92', 'Requests remaining in current window'],
        ['X-RateLimit-Reset', '1743987600', 'Unix timestamp when the hourly window resets'],
        ['X-TokenLimit-Limit', '50000', 'Max tokens allowed per day'],
        ['X-TokenLimit-Remaining', '37550', 'Tokens remaining today'],
        ['X-TokenLimit-Reset', '1744070400', 'Unix timestamp when the daily window resets'],
    ])

    doc.add_heading('Error Response Format', level=1)
    body(doc, 'All errors follow a consistent JSON structure so client code can handle them uniformly. '
         'The error object always contains code, message, status, timestamp, and request_id.', size=10)
    _json_block(doc,
        '{\n'
        '  "error": {\n'
        '    "code": "rate_limit_exceeded",\n'
        '    "message": "You have exceeded your hourly request limit. Try again in 847 seconds.",\n'
        '    "status": 429,\n'
        '    "timestamp": "2026-04-07T14:30:00Z",\n'
        '    "request_id": "mac-req-a1b2c3",\n'
        '    "retry_after": 847\n'
        '  }\n'
        '}')
    styled_table(doc, ['HTTP Status', 'Code', 'When It Happens'],
    [
        ['400', 'bad_request', 'Malformed JSON, invalid field values'],
        ['401', 'authentication_failed', 'Missing/invalid/expired JWT or API key'],
        ['403', 'forbidden', 'Valid auth but insufficient role (e.g. student hitting admin endpoint)'],
        ['404', 'not_found', 'Endpoint or resource does not exist'],
        ['409', 'conflict', 'Duplicate resource (e.g. roll number already registered)'],
        ['413', 'payload_too_large', 'Request body exceeds max size (audio >25MB, image >10MB)'],
        ['422', 'validation_error', 'Schema validation failed — wrong types, missing required fields'],
        ['429', 'rate_limit_exceeded', 'Hourly request limit or daily token quota exceeded'],
        ['500', 'internal_error', 'Unexpected server error — logged and alertable'],
        ['503', 'model_unavailable', 'Requested model is not loaded or all workers are busy'],
    ])

    doc.add_heading('OpenAI SDK Usage Example', level=1)
    body(doc, 'Students can use the standard OpenAI Python SDK with zero modifications — just change base_url:', size=10)
    _json_block(doc,
        'from openai import OpenAI\n\n'
        'client = OpenAI(\n'
        '    base_url="http://mac-server/api/v1",\n'
        '    api_key="mac_sk_live_your_key_here"\n'
        ')\n\n'
        '# Chat\n'
        'response = client.chat.completions.create(\n'
        '    model="auto",\n'
        '    messages=[{"role": "user", "content": "Explain quicksort"}]\n'
        ')\n'
        'print(response.choices[0].message.content)\n\n'
        '# Embeddings\n'
        'emb = client.embeddings.create(\n'
        '    input="Binary search tree",\n'
        '    model="default"\n'
        ')\n'
        'print(len(emb.data[0].embedding))  # 768\n\n'
        '# Speech-to-text (using requests library)\n'
        'import requests\n'
        'resp = requests.post(\n'
        '    "http://mac-server/api/v1/query/speech-to-text",\n'
        '    headers={"Authorization": "Bearer mac_sk_live_xxx"},\n'
        '    files={"audio": open("lecture.mp3", "rb")}\n'
        ')\n'
        'print(resp.json()["text"])')

    save(doc, 'Phase-1-API-Endpoints.docx')


def gen_phase_2():
    """Phase 2 — LLM Models."""
    doc = new_doc('Phase 2 — LLM Models', 'Select, download, feasibility-check, and deploy 5 specialist models')

    doc.add_heading('Model Selection Strategy', level=1)
    body(doc, 'I tested and evaluated dozens of open-source models to find the best specialist for each task category. '
         'The selection criteria were: (1) best benchmark performance in its niche, (2) fits in consumer GPU VRAM with '
         'quantisation, (3) actively maintained with a strong community, and (4) compatible with vLLM for efficient serving.')
    body(doc, 'The result is five models, each the top choice in its domain. Students never need to manually pick — the '
         'API routes automatically based on request content (smart routing).')

    doc.add_heading('The Five Models', level=1)
    styled_table(doc, ['Model ID', 'Name', 'Specialty', 'Params', 'VRAM', 'Quantisation'],
    [
        ['qwen2.5-coder-7b', 'Qwen2.5-Coder 7B', 'Code generation, debugging, explanation', '7B', '~5 GB', 'GPTQ-Int4'],
        ['deepseek-r1-8b', 'DeepSeek-R1 8B', 'Maths, reasoning, step-by-step logic', '8B', '~6 GB', 'AWQ-Int4'],
        ['llava-1.6-7b', 'LLaVA 1.6 7B', 'Image understanding, visual Q&A', '7B', '~8 GB', 'FP16'],
        ['whisper-large-v3', 'Whisper Large v3', 'Speech-to-text, transcription', '1.5B', '~3 GB', 'FP16'],
        ['qwen2.5-14b', 'Qwen2.5 14B', 'General chat, summarisation, writing', '14B', '~10 GB', 'GPTQ-Int4'],
    ])

    doc.add_heading('VRAM Requirements', level=2)
    body(doc, 'Total VRAM for all five models: approximately 32 GB. On a single GPU (e.g., RTX 3060 12GB or RTX 4070 16GB), '
         'models are loaded on demand — only the active model occupies VRAM at any time.')
    add_diagram(doc, dia_model_vram())

    doc.add_heading('Single-PC Feasibility', level=1)
    body(doc, 'For the initial single-PC deployment, models cannot all be loaded simultaneously. The system uses an '
         'intelligent model management strategy:')
    body(doc, '• On-demand loading — When a request comes in for a model that is not in memory, the system loads it '
         'and evicts the least-recently-used model if VRAM is full.')
    body(doc, '• Preloading — The most popular model (Qwen2.5 14B general) stays resident by default. Other models '
         'are loaded as needed.')
    body(doc, '• Configurable timeout — Idle models are automatically unloaded after a configurable period (default: 15 minutes '
         'of no requests) to free VRAM for the next model.')
    body(doc, '• Queue management — If a model is being loaded while new requests arrive for it, those requests are '
         'queued (not rejected) and served once loading completes.')

    doc.add_heading('Smart Routing', level=1)
    body(doc, 'When students send a request with model: "auto", the API gateway inspects the request content and routes '
         'to the best model automatically. This is the recommended way to use the API — students do not need to know '
         'which model handles what.')
    add_diagram(doc, dia_smart_routing())

    styled_table(doc, ['Signal Detected in Request', 'Routed To'],
    [
        ['Code keywords: function, debug, class, python, javascript, etc.', 'qwen2.5-coder-7b'],
        ['Math: solve, prove, equation, integral, derivative, step-by-step', 'deepseek-r1-8b'],
        ['Image file attached in request body', 'llava-1.6-7b'],
        ['Audio file uploaded', 'whisper-large-v3'],
        ['General text: summarise, explain, write, translate, Q&A', 'qwen2.5-14b'],
    ])

    doc.add_heading('Model Management Endpoints — /models', level=1)
    styled_table(doc, ['Method', 'Endpoint', 'Description'],
    [
        ['GET', '/models', 'List all models with status: loaded / downloading / queued / offline'],
        ['GET', '/models/{model_id}', 'Model details — context length, capabilities, benchmarks'],
        ['POST', '/models/{model_id}/load', 'Load model into GPU memory (admin-only)'],
        ['POST', '/models/{model_id}/unload', 'Unload model to free VRAM (admin-only)'],
        ['GET', '/models/{model_id}/health', 'Ping model — latency, memory usage, ready status'],
        ['POST', '/models/download', 'Download a model from HuggingFace (admin-only)'],
        ['GET', '/models/download/{task_id}', 'Check download progress'],
    ])

    save(doc, 'Phase-2-LLM-Models.docx')


def gen_phase_3():
    """Phase 3 — API–Model Integration."""
    doc = new_doc('Phase 3 — API–Model Integration', 'Wire every endpoint to vLLM through LiteLLM proxy')

    doc.add_heading('Integration Architecture', level=1)
    body(doc, 'This phase connects the FastAPI gateway (built in Phase 1) to the LLM models (deployed in Phase 2). '
         'The key middleware component is LiteLLM — an open-source proxy that sits between our API and vLLM workers.')
    body(doc, 'LiteLLM handles three critical responsibilities:')
    body(doc, '1. Request Translation — converts our /query requests into the format each vLLM worker expects.')
    body(doc, '2. Model Routing — based on the smart routing rules, directs each request to the correct model worker.')
    body(doc, '3. Load Balancing — when multiple workers serve the same model, distributes requests using a least-busy strategy.')
    add_diagram(doc, dia_integration_arch())

    doc.add_heading('How a Request Flows', level=2)
    body(doc, 'Here is what happens when a student sends POST /query/chat with model: "auto":')
    body(doc, '1. FastAPI receives the request, validates JWT/API key, checks rate limits.')
    body(doc, '2. The model router analyses the message content and selects the best model (e.g., qwen2.5-coder for code questions).')
    body(doc, '3. FastAPI forwards the request to LiteLLM proxy with the selected model ID.')
    body(doc, '4. LiteLLM looks up the vLLM worker(s) serving that model and picks the least-busy one.')
    body(doc, '5. The vLLM worker runs GPU inference and returns the completion.')
    body(doc, '6. LiteLLM returns the response in OpenAI-compatible format.')
    body(doc, '7. FastAPI logs the usage (tokens in/out, latency) and returns the response to the student.')

    doc.add_heading('Integration Endpoints — /integration', level=1)
    styled_table(doc, ['Method', 'Endpoint', 'Description'],
    [
        ['GET', '/integration/routing-rules', 'Show current routing rules (task type → model)'],
        ['PUT', '/integration/routing-rules', 'Update routing rules (admin-only)'],
        ['GET', '/integration/workers', 'List all vLLM worker nodes and their current load'],
        ['GET', '/integration/workers/{node_id}', 'Single worker — GPU temp, VRAM used, requests in flight'],
        ['POST', '/integration/workers/{node_id}/drain', 'Mark worker as draining (admin-only)'],
        ['GET', '/integration/queue', 'Current global inference queue depth'],
    ])

    doc.add_heading('Scaling Strategy', level=1)
    body(doc, 'The integration layer is designed to scale seamlessly:')
    styled_table(doc, ['Deployment Size', 'Configuration'],
    [
        ['Single PC', 'One vLLM process, models loaded/swapped on demand'],
        ['2–5 PCs', 'Each PC runs a vLLM worker with 1–2 dedicated models; LiteLLM routes by model'],
        ['6–30 PCs', 'Multiple workers per model for redundancy; least-busy routing; auto-failover'],
    ])
    body(doc, 'All worker addresses are stored in configuration (environment variables / LiteLLM config file). '
         'No IP addresses are hardcoded. Adding a new node requires only updating the config and restarting the LiteLLM proxy — '
         'the FastAPI gateway and student code remain completely untouched.')

    doc.add_heading('LiteLLM Configuration', level=2)
    body(doc, 'The LiteLLM proxy is configured with a YAML file that lists all available models, their backends, and '
         'routing strategy. Example structure:')
    config_yaml = (
        'model_list:\n'
        '  - model_name: qwen2.5-coder-7b\n'
        '    litellm_params:\n'
        '      model: openai/qwen2.5-coder-7b\n'
        '      api_base: http://vllm-worker-1:8000/v1\n'
        '  - model_name: deepseek-r1-8b\n'
        '    litellm_params:\n'
        '      model: openai/deepseek-r1-8b\n'
        '      api_base: http://vllm-worker-1:8000/v1\n\n'
        'router_settings:\n'
        '  routing_strategy: least-busy\n'
        '  num_retries: 2\n'
        '  timeout: 120'
    )
    p = doc.add_paragraph()
    r = p.add_run(config_yaml)
    r.font.size = Pt(8)
    r.font.name = 'Courier New'

    save(doc, 'Phase-3-API-Model-Integration.docx')


def gen_phase_4():
    """Phase 4 — API Usage Control."""
    doc = new_doc('Phase 4 — API Usage Control', 'Rate limiting · Token accounting · API key management')

    doc.add_heading('Why Usage Control Matters', level=1)
    body(doc, 'Without usage control, a handful of students running large batch jobs could monopolise the GPU and leave '
         'everyone else waiting indefinitely. Rate limiting and quotas ensure fair, equitable access to the shared compute '
         'resources.')

    doc.add_heading('Request Processing Pipeline', level=1)
    body(doc, 'Every request passes through a multi-stage validation pipeline before reaching the model:')
    add_diagram(doc, dia_rate_limit())
    body(doc, 'At each stage, if a check fails, the request is immediately rejected with the appropriate HTTP error code. '
         'This prevents wasted GPU cycles on unauthorised or over-quota requests.')

    doc.add_heading('API Key Management — /keys', level=1)
    body(doc, 'Every student receives a unique API key upon account creation. I decided to support two types of keys after '
         'considering how students actually use APIs:')
    styled_table(doc, ['Key Type', 'Behaviour', 'Best For'],
    [
        ['Static', 'Never expires; manually rotated by student or admin', 'Quick testing, interactive sessions'],
        ['Refresh', 'Auto-rotates every 30 days; old key has 48-hour grace period', 'Long-running scripts, scheduled jobs'],
    ])
    body(doc, 'Key format: mac_sk_live_<32-character-random-hex>')
    body(doc, 'Security: only the SHA-256 hash of each key is stored in the database. The raw key is shown exactly once — '
         'at creation time. Even database administrators cannot recover a key from its hash.')

    add_diagram(doc, dia_key_lifecycle())

    styled_table(doc, ['Method', 'Endpoint', 'Description'],
    [
        ['GET', '/keys/my-key', 'Get current API key (partially masked) and metadata'],
        ['POST', '/keys/generate', 'Generate new API key (invalidates the old one)'],
        ['GET', '/keys/my-key/stats', 'Tokens consumed against this key — today, week, month'],
        ['DELETE', '/keys/my-key', 'Revoke current key permanently (must re-generate)'],
        ['GET', '/keys/admin/all', 'List all student API keys and status (admin-only)'],
        ['POST', '/keys/admin/revoke', 'Force-revoke a student\'s key (admin-only)'],
    ])

    doc.add_heading('Rate Limiting & Quotas — /quota', level=1)
    body(doc, 'Rate limits are enforced at the Redis layer using a sliding-window algorithm. This gives smooth, fair '
         'rate limiting without the "burst at window edge" problem of fixed-window counters.')

    styled_table(doc, ['Role', 'Daily Token Limit', 'Requests/Hour', 'Max Tokens/Request'],
    [
        ['Student', '50,000', '100', '4,096'],
        ['Faculty', '200,000', '500', '8,192'],
        ['Admin', 'Unlimited', 'Unlimited', '16,384'],
    ])

    styled_table(doc, ['Method', 'Endpoint', 'Description'],
    [
        ['GET', '/quota/limits', 'Show default quota limits per role'],
        ['GET', '/quota/me', 'My personal limits and current consumption'],
        ['PUT', '/quota/admin/user/{roll_no}', 'Override quota for a specific user (admin-only)'],
        ['GET', '/quota/admin/exceeded', 'List users who exceeded their quota today (admin-only)'],
    ])

    doc.add_heading('Rate Limit Response Headers', level=2)
    body(doc, 'Every API response includes standard rate limit headers so client code can self-throttle:')
    headers_text = (
        'X-RateLimit-Limit: 100\n'
        'X-RateLimit-Remaining: 92\n'
        'X-RateLimit-Reset: 1743987600\n'
        'X-TokenLimit-Limit: 50000\n'
        'X-TokenLimit-Remaining: 37550\n'
        'X-TokenLimit-Reset: 1744070400'
    )
    p = doc.add_paragraph()
    r = p.add_run(headers_text)
    r.font.size = Pt(9)
    r.font.name = 'Courier New'
    body(doc, 'When limits are exceeded, the API returns HTTP 429 with a retry_after field in seconds.')

    save(doc, 'Phase-4-API-Usage-Control.docx')


def gen_phase_5():
    """Phase 5 — Web Interface."""
    doc = new_doc('Phase 5 — Web Interface', 'Dashboard · User Management · Admin Panel · Playground')

    doc.add_heading('Purpose', level=1)
    body(doc, 'Not every student wants to write code to use the AI platform. The web interface provides a visual, '
         'point-and-click experience where students can chat with models, track their usage, manage API keys, and '
         'explore available models — all from a browser.')
    body(doc, 'The frontend is built with React and Tailwind CSS, served by the FastAPI backend. It communicates '
         'with the same REST API that SDK users call, so there is a single source of truth.')

    doc.add_heading('Dashboard Layout', level=1)
    body(doc, 'The student dashboard is the main landing page after login. It shows usage at a glance, available models, '
         'quick-start code snippets, and a weekly usage trend chart.')
    add_diagram(doc, dia_dashboard())

    doc.add_heading('Student-Facing Pages', level=2)
    styled_table(doc, ['Page', 'Endpoint', 'What It Shows'],
    [
        ['Dashboard', '/ui/dashboard', 'Usage summary cards, model status, quick-start code, usage chart'],
        ['API Keys', '/ui/keys', 'View key (masked), copy to clipboard, regenerate, key type toggle'],
        ['History', '/ui/history', 'Request history table — timestamp, model, tokens, latency, status'],
        ['Playground', '/ui/playground', 'Interactive chat — select model, send messages, see streamed responses'],
    ])

    doc.add_heading('Admin Panel', level=1)
    body(doc, 'Administrators (faculty and lab staff) get an extended interface for managing users, models, and monitoring '
         'the platform.')
    styled_table(doc, ['Page', 'Endpoint', 'What It Shows'],
    [
        ['User Management', '/ui/admin/users', 'Full user list with roles, quotas, last activity, status'],
        ['Create Users', '/ui/admin/users/create', 'Bulk-create students from CSV (roll no, name, department)'],
        ['Edit User', '/ui/admin/users/{roll}', 'Change role, override quotas, restrict model access'],
        ['Model Management', '/ui/admin/models', 'Load/unload models, see which node serves what, GPU stats'],
        ['System Logs', '/ui/admin/logs', 'Live request logs, error rates, latency percentiles, throughput'],
        ['Analytics', '/ui/admin/analytics', 'Daily active users, peak hours, top models, monthly trends'],
    ])

    doc.add_heading('User Management Flow', level=2)
    body(doc, 'The admin can create users individually or in bulk via CSV upload. The CSV format is simple: '
         'roll_number, name, department. On upload, the system generates a default password for each student '
         '(which they must change on first login) and creates their initial API key.')
    body(doc, 'Admins can assign three roles:')
    body(doc, '• Student — standard access with default quotas')
    body(doc, '• Faculty — elevated quotas, can view their students\' usage')
    body(doc, '• Admin — full access to all management features')

    save(doc, 'Phase-5-Web-Interface.docx')


def gen_phase_6():
    """Phase 6 — Guardrails."""
    doc = new_doc('Phase 6 — Guardrails', 'Input + Output Filtering · Safety Checks')

    doc.add_heading('Why Guardrails?', level=1)
    body(doc, 'Running open-source LLMs in an academic environment requires safety measures. Guardrails protect both '
         'the students (from harmful AI output) and the institution (from misuse). They act as filters on both the input '
         'side (what users send to models) and the output side (what models send back).')

    doc.add_heading('Filtering Pipeline', level=1)
    body(doc, 'Every request and response passes through the guardrails pipeline:')
    add_diagram(doc, dia_guardrails())

    doc.add_heading('Input Checks', level=2)
    styled_table(doc, ['Check', 'Action', 'Description'],
    [
        ['Prompt Injection Detection', 'Block + log', 'Detects attempts to override system prompts or jailbreak the model'],
        ['Blocked Topic Detection', 'Block', 'Filters requests about violence, illegal activities, self-harm'],
        ['Max Prompt Length', 'Reject', 'Configurable per role; prevents resource abuse via extremely long prompts'],
        ['Academic Integrity Check', 'Flag + disclaimer', 'Detects full assignment/essay requests; adds integrity notice'],
    ])

    doc.add_heading('Output Checks', level=2)
    styled_table(doc, ['Check', 'Action', 'Description'],
    [
        ['PII Redaction', 'Redact', 'Strips emails, phone numbers, addresses from model output'],
        ['Harmful Content Detection', 'Block', 'Catches harmful content the model may generate despite input filtering'],
        ['Academic Integrity Notice', 'Append', 'Adds a disclaimer when output appears to be complete coursework'],
        ['Source Attribution', 'Append', 'For RAG responses, ensures source citations are included'],
    ])

    doc.add_heading('Guardrails Endpoints — /guardrails', level=1)
    styled_table(doc, ['Method', 'Endpoint', 'Description'],
    [
        ['POST', '/guardrails/check-input', 'Run text through input content filter'],
        ['POST', '/guardrails/check-output', 'Run model output through safety filter'],
        ['GET', '/guardrails/rules', 'List active guardrail rules (admin-only)'],
        ['PUT', '/guardrails/rules', 'Update rules — blocked categories, max length (admin-only)'],
    ])

    doc.add_heading('Implementation Approach', level=2)
    body(doc, 'The guardrails system uses a layered approach:')
    body(doc, '1. Keyword-based filters — fast, rule-based checks for known bad patterns (runs first).')
    body(doc, '2. Classifier-based filters — a small, fast classifier model (distilbert-based) for nuanced content '
         'categorisation (runs second, only if keyword filters pass).')
    body(doc, '3. Regex-based PII detection — pattern matching for emails, phone numbers, and structured data in outputs.')
    body(doc, 'All rules are stored in the database and configurable through the admin panel — no code changes '
         'needed to add or modify rules.')

    save(doc, 'Phase-6-Guardrails.docx')


def gen_phase_7():
    """Phase 7 — Knowledgebase + RAG."""
    doc = new_doc('Phase 7 — Knowledgebase + RAG', 'Vector DB · Document Ingestion · Retrieval-Augmented Generation')

    doc.add_heading('What is RAG and Why?', level=1)
    body(doc, 'Retrieval-Augmented Generation (RAG) lets us ground the AI\'s answers in actual course material. Instead of '
         'relying solely on what the model learned during training (which may be outdated or inaccurate), RAG retrieves '
         'relevant passages from uploaded textbooks and includes them as context when generating a response.')
    body(doc, 'This means students can ask "Explain the difference between process and thread" and get an answer that '
         'directly references their OS textbook (Galvin, Chapter 3) — with page citations.')

    doc.add_heading('RAG Pipeline', level=1)
    body(doc, 'The pipeline has two flows: document ingestion (one-time, per document) and query retrieval (per student question).')
    add_diagram(doc, dia_rag_pipeline())

    doc.add_heading('Document Ingestion Flow', level=2)
    body(doc, '1. Admin uploads a PDF, DOCX, or TXT file through the admin panel or API.')
    body(doc, '2. The document is split into chunks of ~512 tokens with 50-token overlap between chunks '
         '(this ensures no information is lost at chunk boundaries).')
    body(doc, '3. Each chunk is converted into a 768-dimensional vector embedding using an embedding model.')
    body(doc, '4. The embedding and chunk text are stored in Qdrant, our vector database, along with metadata '
         '(document title, page number, chapter).')

    doc.add_heading('Query Retrieval Flow', level=2)
    body(doc, '1. Student sends a question to POST /rag/query.')
    body(doc, '2. The question is embedded using the same embedding model.')
    body(doc, '3. Qdrant performs a similarity search, returning the top-k (default 5) most relevant chunks.')
    body(doc, '4. The retrieved chunks are concatenated with the original question and sent to the LLM.')
    body(doc, '5. The LLM generates an answer grounded in the retrieved content.')
    body(doc, '6. The response includes source citations (document, chapter, page, relevance score).')

    doc.add_heading('RAG Endpoints — /rag', level=1)
    styled_table(doc, ['Method', 'Endpoint', 'Description'],
    [
        ['POST', '/rag/ingest', 'Upload PDF/DOCX/TXT — chunk, embed, store in Qdrant'],
        ['GET', '/rag/documents', 'List all ingested documents with metadata'],
        ['GET', '/rag/documents/{id}', 'Single document details and its chunks'],
        ['DELETE', '/rag/documents/{id}', 'Remove document from knowledgebase (admin-only)'],
        ['POST', '/rag/query', 'Ask a question — retrieve chunks → LLM → cited answer'],
        ['GET', '/rag/query/{id}/sources', 'Get source citations for a RAG response'],
        ['POST', '/rag/collections', 'Create a named collection: DSA, DBMS, OS, etc. (admin-only)'],
        ['GET', '/rag/collections', 'List all collections'],
    ])

    doc.add_heading('Collections', level=2)
    body(doc, 'Documents are organised into collections by subject (e.g., "DSA", "DBMS", "Operating Systems"). '
         'When querying, students can optionally specify a collection to limit the search scope — this improves '
         'both relevance and speed.')

    doc.add_heading('Technical Details', level=2)
    styled_table(doc, ['Parameter', 'Value', 'Rationale'],
    [
        ['Chunk size', '512 tokens', 'Balances context richness with retrieval precision'],
        ['Chunk overlap', '50 tokens', 'Prevents information loss at boundaries'],
        ['Embedding dimensions', '768', 'Standard dimension for modern embedding models'],
        ['Top-k retrieval', '5 (default)', 'Good balance of context coverage without overwhelming the LLM'],
        ['Vector DB', 'Qdrant', 'Purpose-built for similarity search, supports filtering by collection'],
    ])

    save(doc, 'Phase-7-Knowledgebase-RAG.docx')


def gen_phase_8():
    """Phase 8 — Retrieval + Search."""
    doc = new_doc('Phase 8 — Retrieval + Search', 'SearXNG web search · Wikipedia · Real-time grounded answers')

    doc.add_heading('Purpose', level=1)
    body(doc, 'The knowledgebase (Phase 7) covers course material, but students also need answers about current events, '
         'recent research, real-time data, and topics beyond our textbook collection. This phase adds web search '
         'capability through SearXNG — a self-hosted, privacy-respecting meta-search engine.')
    body(doc, 'The key feature is grounded search — the AI does not just return raw search results. It reads the '
         'search results, synthesises an answer, and cites its sources. Students get a coherent response backed '
         'by real web sources.')

    doc.add_heading('Search Pipeline', level=1)
    add_diagram(doc, dia_search_flow())

    doc.add_heading('How It Works', level=2)
    body(doc, '1. Student sends a query to POST /search/grounded.')
    body(doc, '2. The server queries SearXNG, which aggregates results from Google, Bing, DuckDuckGo, and Wikipedia.')
    body(doc, '3. The top results (titles, URLs, snippets) are collected.')
    body(doc, '4. These results, along with the original question, are sent to the LLM as context.')
    body(doc, '5. The LLM generates a comprehensive answer with inline citations.')
    body(doc, '6. The response includes the source URLs so students can verify independently.')

    doc.add_heading('Search Endpoints — /search', level=1)
    styled_table(doc, ['Method', 'Endpoint', 'Description'],
    [
        ['POST', '/search/web', 'Query SearXNG — returns top results from Google, Bing, DuckDuckGo, Wikipedia'],
        ['POST', '/search/wikipedia', 'Targeted Wikipedia search with summary extraction'],
        ['POST', '/search/grounded', 'Search + LLM — retrieves web results, generates cited answer'],
        ['GET', '/search/cache', 'List recently cached search results (reduces repeated fetches)'],
    ])

    doc.add_heading('Caching Strategy', level=2)
    body(doc, 'Search results are cached in Redis for a configurable TTL (default: 1 hour). This means:')
    body(doc, '• If multiple students ask similar questions within an hour, only one SearXNG query is made.')
    body(doc, '• Reduces load on external search engines and speeds up response times.')
    body(doc, '• Cache keys are normalised (lowercased, stripped) so minor variations in wording still hit the cache.')
    body(doc, '• Admins can view and clear the cache through the /search/cache endpoint.')

    doc.add_heading('Why SearXNG?', level=2)
    body(doc, 'I chose SearXNG over commercial search APIs for several reasons:')
    body(doc, '• Self-hosted — runs in a Docker container alongside our stack. No external API keys needed.')
    body(doc, '• Privacy — search queries stay within our network. No data sent to third parties.')
    body(doc, '• Meta-search — aggregates results from multiple search engines for better coverage.')
    body(doc, '• Free — no per-query charges, which matters when 500+ students are searching.')
    body(doc, '• Configurable — we can enable/disable specific search engines and set result limits.')

    save(doc, 'Phase-8-Retrieval-Search.docx')


# ════════════════════════════════════════════════════════════════
#  MAIN
# ════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    print(f'\nGenerating MAC documentation in: {OUT}\n')
    gen_phase_0()
    gen_phase_1()
    gen_phase_2()
    gen_phase_3()
    gen_phase_4()
    gen_phase_5()
    gen_phase_6()
    gen_phase_7()
    gen_phase_8()
    print(f'\nDone — 9 documents generated.\n')
