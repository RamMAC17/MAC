"""
MAC (MBM AI Cloud) — Comprehensive Project Report
Generates a professional Word document covering the entire platform.

Run:  python docs/generate_final_report.py
Output: docs/MAC-Final-Project-Report.docx
"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
from pathlib import Path
import textwrap

OUT = Path(__file__).parent
plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['font.sans-serif'] = ['Segoe UI', 'Arial', 'Helvetica', 'DejaVu Sans']

# ── Colour palette ────────────────────────────────────────────
C_DARK    = '#0d1b2a'
C_PRIMARY = '#1b4965'
C_SEC     = '#2a6f97'
C_ACCENT  = '#468faf'
C_LIGHT   = '#a9d6e5'
C_PALE    = '#e8f4f8'
C_ORANGE  = '#e76f51'
C_GREEN   = '#52b788'
C_TEAL    = '#2ec4b2'
C_AMBER   = '#f4a261'
C_RED     = '#e63946'
C_GRAY    = '#6c757d'
C_LGRAY   = '#f6f8fa'

# ════════════════════════════════════════════════════════════════
#  HELPERS
# ════════════════════════════════════════════════════════════════

def _shade(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex.lstrip('#'))
    shading.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading)

def _set_col_widths(table, widths_inches):
    for row in table.rows:
        for i, w in enumerate(widths_inches):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)

def _set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), val.get('val', 'single'))
        el.set(qn('w:sz'), val.get('sz', '4'))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), val.get('color', '000000'))
        borders.append(el)
    tcPr.append(borders)

def new_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.54)
        sec.right_margin = Cm(2.54)
    return doc

def body(doc, text, bold=False, italic=False, size=11, justify=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(6)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1b, 0x49, 0x65)
    return h

def styled_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        _shade(c, C_PRIMARY)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            if ri % 2 == 1:
                _shade(c, C_LGRAY)
    if col_widths:
        _set_col_widths(t, col_widths)
    doc.add_paragraph()
    return t

def code_block(doc, text, size=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Consolas'
    return p

def add_diagram(doc, buf, width=6.2):
    doc.add_picture(buf, width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def bullet(doc, text, bold_prefix=None, size=11):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(size)
        r1.font.name = 'Calibri'
        r2 = p.add_run(text)
        r2.font.size = Pt(size)
        r2.font.name = 'Calibri'
    else:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

# ════════════════════════════════════════════════════════════════
#  DIAGRAM PRIMITIVES
# ════════════════════════════════════════════════════════════════

def _box(ax, x, y, w, h, label, color=C_PRIMARY, tc='white', fs=9, fw='bold', ls=None):
    b = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.03",
                       facecolor=color, edgecolor='#dddddd', linewidth=0.6)
    ax.add_patch(b)
    lines = label.split('\n')
    if len(lines) == 1:
        ax.text(x + w/2, y + h/2, label, ha='center', va='center',
                fontsize=fs, color=tc, fontweight=fw)
    else:
        top = lines[0]
        bot = '\n'.join(lines[1:])
        ax.text(x + w/2, y + h*0.62, top, ha='center', va='center',
                fontsize=fs, color=tc, fontweight=fw)
        ax.text(x + w/2, y + h*0.30, bot, ha='center', va='center',
                fontsize=fs - 1.5, color=tc, fontweight='normal', alpha=0.85)

def _arrow(ax, x1, y1, x2, y2, color='#555555', lw=1.5, style='->'):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle=style, color=color, lw=lw))

def _fig(w=10, h=6, xl=(0, 10), yl=(0, 6)):
    fig, ax = plt.subplots(figsize=(w, h))
    ax.set_xlim(xl)
    ax.set_ylim(yl)
    ax.axis('off')
    fig.patch.set_facecolor('white')
    return fig, ax

def _to_buf(fig):
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    return buf

# ════════════════════════════════════════════════════════════════
#  DIAGRAMS
# ════════════════════════════════════════════════════════════════

def dia_architecture():
    fig, ax = _fig(11, 9, (0, 11), (0, 9))
    ax.text(5.5, 8.6, 'MAC — System Architecture', ha='center', fontsize=14,
            fontweight='bold', color=C_DARK)

    # Users
    _box(ax, 2.5, 7.5, 6, 0.7, 'Students / Faculty — College LAN (PWA Frontend)', C_TEAL, 'white', 10)
    _arrow(ax, 5.5, 7.5, 5.5, 7.0)

    # Nginx
    _box(ax, 3, 6.2, 5, 0.7, 'Nginx — Reverse Proxy · TLS · SSE Streaming', C_ORANGE, 'white', 10)
    _arrow(ax, 5.5, 6.2, 5.5, 5.7)

    # FastAPI
    _box(ax, 1.2, 4.9, 8.6, 0.7, 'FastAPI Gateway — Auth · Rate Limiting · Guardrails · Routing · 50+ Endpoints', C_PRIMARY, 'white', 10)

    # Arrows to services
    _arrow(ax, 2.5, 4.9, 1.5, 4.3)
    _arrow(ax, 5.5, 4.9, 5.5, 4.3)
    _arrow(ax, 8.5, 4.9, 9.5, 4.3)

    # Services row
    _box(ax, 0.2, 3.5, 2.6, 0.7, 'PostgreSQL / SQLite\nUsers · Logs · Quotas', C_SEC, 'white', 8)
    _box(ax, 4, 3.5, 3, 0.7, 'LiteLLM Proxy\nRouting · Load Balance', C_SEC, 'white', 8)
    _box(ax, 8.2, 3.5, 2.6, 0.7, 'Qdrant · Redis\nVectors · Cache', C_SEC, 'white', 8)

    _arrow(ax, 5.5, 3.5, 5.5, 3.0)

    # vLLM / Ollama
    _box(ax, 2.5, 2.2, 6, 0.7, 'vLLM / Ollama — GPU Inference Engine', C_DARK, 'white', 10)

    # GPU Nodes
    _box(ax, 1.0, 0.8, 2.2, 0.6, 'GPU Node 1\nRTX 3060 12GB', C_ACCENT, 'white', 7)
    _box(ax, 4.4, 0.8, 2.2, 0.6, 'GPU Node 2', C_ACCENT, 'white', 8)
    _box(ax, 7.8, 0.8, 2.2, 0.6, 'GPU Node N', C_ACCENT, 'white', 8)
    ax.text(6.95, 1.05, '...', fontsize=16, color=C_GRAY, ha='center')
    _arrow(ax, 3.5, 2.2, 2.1, 1.45)
    _arrow(ax, 5.5, 2.2, 5.5, 1.45)
    _arrow(ax, 7.5, 2.2, 8.9, 1.45)

    # SearXNG side
    _box(ax, 0.2, 2.2, 2, 0.6, 'SearXNG\nWeb Search', C_AMBER, 'white', 8)
    _arrow(ax, 1.2, 4.9, 1.2, 2.85, color=C_AMBER, lw=1)

    return _to_buf(fig)


def dia_roadmap():
    fig, ax = _fig(12, 5.5, (0, 12), (0, 5.5))
    phases = [
        ('Phase 1', 'API Endpoints'),
        ('Phase 2', 'LLM Models'),
        ('Phase 3', 'API–Model\nIntegration'),
        ('Phase 4', 'Usage Control'),
        ('Phase 5', 'Web Interface'),
        ('Phase 6', 'Guardrails'),
        ('Phase 7', 'Knowledge\nBase + RAG'),
        ('Phase 8', 'Retrieval\n+ Search'),
    ]
    colors = ['#0d1b2a', '#1b3a4b', '#1b4965', '#2a6f97',
              '#468faf', '#61a5c2', '#52b788', '#e76f51']
    tcs = ['white'] * 8
    xs = [0.3, 3.1, 5.9, 8.7]
    bw, bh = 2.4, 1.1
    for i, (label, sub) in enumerate(phases):
        row = 0 if i < 4 else 1
        col = i % 4
        x = xs[col]
        y = 3.5 if row == 0 else 1.5
        _box(ax, x, y, bw, bh, f'{label}\n{sub}', colors[i], tcs[i], fs=10)
    for i in range(3):
        _arrow(ax, xs[i] + bw, 3.5 + bh/2, xs[i+1], 3.5 + bh/2)
    ax.annotate('', xy=(xs[0] + bw/2, 1.5 + bh), xytext=(xs[3] + bw/2, 3.5),
                arrowprops=dict(arrowstyle='->', color='#555', lw=1.5,
                                connectionstyle='arc3,rad=0.4'))
    for i in range(3):
        _arrow(ax, xs[i] + bw, 1.5 + bh/2, xs[i+1], 1.5 + bh/2)
    ax.text(6, 5.1, 'MAC — 8-Phase Build Roadmap', ha='center', fontsize=14,
            fontweight='bold', color=C_DARK)
    return _to_buf(fig)


def dia_auth_flow():
    fig, ax = _fig(12, 5, (0, 12), (0, 5))
    ax.text(6, 4.7, 'Unified Authentication Flow', ha='center', fontsize=13,
            fontweight='bold', color=C_DARK)
    steps = [
        (0.1, 'Student\nOpens PWA', C_TEAL),
        (2.0, 'Enter\nRoll No.\n+ DOB', C_ACCENT),
        (3.9, 'POST\n/auth/verify', C_SEC),
        (5.8, 'Check\nRegistry\n(DB Match)', C_PRIMARY),
        (7.7, 'Create User\nor Return\nExisting', C_GREEN),
        (9.8, 'JWT Tokens\n+ Profile', C_DARK),
    ]
    bw, bh = 1.7, 1.2
    y = 2.0
    for x, label, color in steps:
        _box(ax, x, y, bw, bh, label, color, 'white', 8)
    for i in range(len(steps) - 1):
        _arrow(ax, steps[i][0] + bw, y + bh/2, steps[i+1][0], y + bh/2)

    # Set password branch
    ax.text(10, 1.3, 'First login?\nmust_change_password=true', fontsize=7, color=C_AMBER,
            ha='center', fontstyle='italic')
    _arrow(ax, 10.65, 2.0, 10.65, 1.55, color=C_AMBER, lw=1.2)
    _box(ax, 9.4, 0.3, 2.5, 0.7, 'Set Password\nPOST /auth/set-password', C_AMBER, 'white', 7)

    # DOB mismatch
    ax.text(6.65, 1.4, 'DOB mismatch → 401', fontsize=7, color=C_RED,
            ha='center', fontweight='bold')
    _arrow(ax, 6.65, 2.0, 6.65, 1.6, color=C_RED, lw=1)

    return _to_buf(fig)


def dia_module_structure():
    fig, ax = _fig(12, 8, (0, 12), (0, 8))
    ax.text(6, 7.7, 'MAC — Module Architecture', ha='center', fontsize=14,
            fontweight='bold', color=C_DARK)

    # Routers
    _box(ax, 0.3, 5.8, 11.4, 1.5, '', C_PALE, C_DARK, 9)
    ax.text(6, 7.0, 'Routers (API Layer) — 11 Modules', fontsize=10, fontweight='bold',
            color=C_PRIMARY, ha='center')
    routers = ['auth', 'explore', 'query', 'usage', 'models', 'keys',
               'quota', 'guardrails', 'rag', 'search', 'integration']
    for i, r in enumerate(routers):
        col = i % 6
        row = i // 6
        x = 0.6 + col * 1.9
        y = 6.35 - row * 0.55
        _box(ax, x, y, 1.7, 0.45, r, C_PRIMARY, 'white', 7)

    # Services
    _box(ax, 0.3, 3.8, 11.4, 1.5, '', '#f0f4f8', C_DARK, 9)
    ax.text(6, 5.0, 'Services (Business Logic) — 7 Modules', fontsize=10, fontweight='bold',
            color=C_SEC, ha='center')
    services = ['auth_service', 'llm_service', 'model_service', 'guardrail_service',
                'rag_service', 'search_service', 'usage_service']
    for i, s in enumerate(services):
        x = 0.5 + i * 1.6
        _box(ax, x, 4.1, 1.45, 0.5, s.replace('_', '\n'), C_SEC, 'white', 6.5)

    # Middleware
    _box(ax, 0.3, 2.5, 5.2, 0.9, '', '#fff3e0', C_DARK, 9)
    ax.text(2.9, 3.15, 'Middleware', fontsize=10, fontweight='bold', color=C_ORANGE, ha='center')
    _box(ax, 0.6, 2.65, 2.2, 0.5, 'Auth\nMiddleware', C_ORANGE, 'white', 7)
    _box(ax, 3.1, 2.65, 2.2, 0.5, 'Rate\nLimiter', C_ORANGE, 'white', 7)

    # Models (ORM)
    _box(ax, 6.0, 2.5, 5.7, 0.9, '', '#e8f5e9', C_DARK, 9)
    ax.text(8.85, 3.15, 'Database Models (ORM)', fontsize=10, fontweight='bold', color=C_GREEN, ha='center')
    orm = ['User', 'StudentRegistry', 'UsageLog', 'GuardrailRule', 'QuotaOverride', 'RAG*']
    for i, m in enumerate(orm):
        x = 6.2 + i * 0.9
        _box(ax, x, 2.65, 0.8, 0.5, m, C_GREEN, 'white', 6)

    # Schemas
    _box(ax, 0.3, 1.2, 5.2, 0.9, '', '#e3f2fd', C_DARK, 9)
    ax.text(2.9, 1.85, 'Pydantic Schemas — 11 Modules', fontsize=9, fontweight='bold',
            color=C_ACCENT, ha='center')
    ax.text(2.9, 1.45, 'auth · chat · explore · guardrails · integration\n'
            'keys · models · quota · rag · search · usage',
            fontsize=7, color=C_DARK, ha='center', linespacing=1.4)

    # Utils
    _box(ax, 6.0, 1.2, 5.7, 0.9, '', '#fce4ec', C_DARK, 9)
    ax.text(8.85, 1.85, 'Utils & Config', fontsize=10, fontweight='bold', color=C_RED, ha='center')
    ax.text(8.85, 1.45, 'security.py · config.py · database.py\n'
            'JWT · bcrypt · SQLAlchemy async engine',
            fontsize=7, color=C_DARK, ha='center', linespacing=1.4)

    # External
    _box(ax, 0.3, 0.1, 11.4, 0.7, '', '#e0e0e0', C_DARK, 9)
    ax.text(6, 0.45, 'External: PostgreSQL/SQLite · Redis · Qdrant · SearXNG · vLLM/Ollama · LiteLLM · Nginx',
            fontsize=9, fontweight='bold', color=C_DARK, ha='center')

    # Arrows
    _arrow(ax, 6, 5.8, 6, 5.35, color=C_GRAY, lw=1.5)
    _arrow(ax, 2.9, 3.8, 2.9, 3.45, color=C_GRAY, lw=1.5)
    _arrow(ax, 8.85, 3.8, 8.85, 3.45, color=C_GRAY, lw=1.5)

    return _to_buf(fig)


def dia_guardrail_pipeline():
    fig, ax = _fig(12, 5, (0, 12), (0, 5))
    ax.text(6, 4.7, 'Guardrails — Input & Output Filtering Pipeline', ha='center',
            fontsize=13, fontweight='bold', color=C_DARK)
    y = 2.0
    bw, bh = 1.8, 0.9
    pipeline = [
        (0.1, 'User\nInput', C_TEAL),
        (2.1, 'Prompt\nInjection\nDetect', C_RED),
        (4.1, 'Harmful\nContent\nFilter', C_ORANGE),
        (6.1, 'LLM\nInference', C_PRIMARY),
        (8.1, 'PII\nRedaction\nFilter', C_AMBER),
        (10.1, 'Safe\nResponse', C_GREEN),
    ]
    for x, label, color in pipeline:
        _box(ax, x, y, bw, bh, label, color, 'white', 8)
    for i in range(len(pipeline) - 1):
        _arrow(ax, pipeline[i][0] + bw, y + bh/2, pipeline[i+1][0], y + bh/2)
    checks_in = ['Jailbreak / override attempts', 'Academic dishonesty keywords', 'Max prompt length (32K chars)']
    for i, txt in enumerate(checks_in):
        ax.text(2.1, 1.5 - i*0.35, f'• {txt}', fontsize=7.5, color=C_DARK)
    ax.text(2.1, 1.5 - 3*0.35 + 0.1, 'INPUT CHECKS', fontsize=7, color=C_RED, fontweight='bold')
    checks_out = ['Email redaction', 'Phone number redaction', 'ID/card number masking']
    for i, txt in enumerate(checks_out):
        ax.text(8.1, 1.5 - i*0.35, f'• {txt}', fontsize=7.5, color=C_DARK)
    ax.text(8.1, 1.5 - 3*0.35 + 0.1, 'OUTPUT CHECKS', fontsize=7, color=C_AMBER, fontweight='bold')
    return _to_buf(fig)


def dia_rag_pipeline():
    fig, ax = _fig(12, 7, (0, 12), (0, 7))
    ax.text(6, 6.7, 'RAG Pipeline — Ingestion & Query', ha='center',
            fontsize=13, fontweight='bold', color=C_DARK)
    ax.text(1, 5.9, 'Document Ingestion', fontsize=10, fontweight='bold', color=C_SEC)
    ing = [
        (0.2, 'Upload\nPDF/DOCX/TXT', C_TEAL),
        (2.5, 'Chunking\n512 tokens\noverlap 50', C_SEC),
        (4.8, 'Embedding\n768-dim vectors', C_ACCENT),
        (7.1, 'Store in\nQdrant', C_PRIMARY),
    ]
    bw, bh = 2, 0.8
    y = 4.8
    for x, label, color in ing:
        _box(ax, x, y, bw, bh, label, color, 'white', 9)
    for i in range(len(ing) - 1):
        _arrow(ax, ing[i][0] + bw, y + bh/2, ing[i+1][0], y + bh/2)
    ax.plot([0.5, 11.5], [4.3, 4.3], color='#ddd', lw=1, ls='--')
    ax.text(1, 4.0, 'Query & Retrieval', fontsize=10, fontweight='bold', color=C_SEC)
    q_top = [
        (0.2, 'Student\nQuestion', C_TEAL),
        (2.5, 'Embed\nQuery', C_SEC),
        (4.8, 'Similarity\nSearch (top-k)', C_ACCENT),
    ]
    y2 = 3.0
    for x, label, color in q_top:
        _box(ax, x, y2, bw, bh, label, color, 'white', 9)
    for i in range(len(q_top) - 1):
        _arrow(ax, q_top[i][0] + bw, y2 + bh/2, q_top[i+1][0], y2 + bh/2)
    _arrow(ax, 6.8, y2 + bh/2, 7.1, y2 + bh/2)
    _box(ax, 7.1, 2.5, 2.5, 1.3, 'LLM\nRetrieved Chunks\n+ Question', C_PRIMARY, 'white', 9)
    _arrow(ax, 9.6, 3.15, 10, 3.15)
    _box(ax, 9.7, 1.4, 2, 0.8, 'Answer with\nCitations', C_GREEN, 'white', 9)
    ax.annotate('', xy=(5.8, 3.8), xytext=(8.1, 4.8),
                arrowprops=dict(arrowstyle='->', color=C_AMBER, lw=1.5, ls='--'))
    ax.text(7.6, 4.35, 'vector index', fontsize=7, color=C_AMBER, fontstyle='italic', rotation=-30)
    return _to_buf(fig)


def dia_request_pipeline():
    fig, ax = _fig(12, 4.5, (0, 12), (0, 4.5))
    ax.text(6, 4.1, 'Request Processing Pipeline', ha='center', fontsize=13,
            fontweight='bold', color=C_DARK)
    steps = [
        (0.1, 'Incoming\nRequest', C_TEAL),
        (2.0, 'JWT / API\nKey Verify', C_SEC),
        (3.9, 'Rate\nLimit\nCheck', C_ACCENT),
        (5.8, 'Token\nQuota\nCheck', C_PRIMARY),
        (7.7, 'Guardrail\nInput\nFilter', C_ORANGE),
        (9.8, 'LLM\nInference', C_GREEN),
    ]
    bw, bh = 1.7, 1.1
    y = 1.5
    for x, label, color in steps:
        _box(ax, x, y, bw, bh, label, color, 'white', 8)
    for i in range(len(steps) - 1):
        _arrow(ax, steps[i][0] + bw, y + bh/2, steps[i+1][0], y + bh/2)
    for i, label in [(1, '401'), (2, '429'), (3, '429'), (4, '400')]:
        ax.annotate('', xy=(steps[i][0] + bw/2, y - 0.15),
                    xytext=(steps[i][0] + bw/2, y),
                    arrowprops=dict(arrowstyle='->', color=C_RED, lw=1.2))
        ax.text(steps[i][0] + bw/2, y - 0.4, f'Reject {label}', ha='center',
                fontsize=7, color=C_RED, fontweight='bold')
    return _to_buf(fig)


def dia_db_schema():
    fig, ax = _fig(12, 7, (0, 12), (0, 7))
    ax.text(6, 6.7, 'Database Schema (Entity-Relationship)', ha='center',
            fontsize=13, fontweight='bold', color=C_DARK)

    tables = [
        (0.2, 4.0, 'student_registry', ['id (PK)', 'roll_number (UQ)', 'name', 'department', 'dob', 'batch_year']),
        (4.2, 4.0, 'users', ['id (PK)', 'roll_number (UQ)', 'name', 'email', 'department',
                               'role', 'password_hash', 'must_change_password',
                               'is_active', 'api_key (UQ)', 'created_at']),
        (8.5, 4.8, 'refresh_tokens', ['id (PK)', 'user_id (FK)', 'token_hash', 'expires_at', 'revoked']),
        (8.5, 2.6, 'usage_logs', ['id (PK)', 'user_id (FK)', 'model', 'endpoint',
                                    'tokens_in', 'tokens_out', 'latency_ms']),
        (0.2, 1.0, 'quota_overrides', ['id (PK)', 'user_id (FK, UQ)', 'daily_tokens',
                                         'requests_per_hour', 'reason']),
        (4.2, 1.0, 'guardrail_rules', ['id (PK)', 'category', 'action', 'pattern',
                                          'enabled', 'priority']),
        (8.5, 0.5, 'rag_collections\n+ rag_documents', ['collection: id, name, description',
                                                           'document: id, collection_id (FK)',
                                                           'title, chunk_count, status']),
    ]
    for x, y_base, title, fields in tables:
        w = 3.8 if 'users' in title else 3.6 if 'rag_' in title else 3.2
        h = 0.18 * len(fields) + 0.5
        _box(ax, x, y_base, w, h, '', '#f0f4f8', C_DARK, 9)
        ax.text(x + w/2, y_base + h - 0.2, title, fontsize=8, fontweight='bold',
                color=C_PRIMARY, ha='center')
        for i, f in enumerate(fields):
            ax.text(x + 0.15, y_base + h - 0.42 - i*0.18, f, fontsize=6, color=C_DARK,
                    family='monospace')

    # Relations
    _arrow(ax, 3.4, 5.0, 4.2, 5.0, color=C_SEC, lw=1.5)
    _arrow(ax, 8.0, 5.2, 8.5, 5.2, color=C_SEC, lw=1.5)
    _arrow(ax, 8.0, 4.5, 8.5, 3.6, color=C_SEC, lw=1.5)
    _arrow(ax, 5.5, 4.0, 1.8, 2.2, color=C_SEC, lw=1)

    return _to_buf(fig)


def dia_deployment():
    fig, ax = _fig(11, 6, (0, 11), (0, 6))
    ax.text(5.5, 5.7, 'Docker Compose — Deployment Architecture', ha='center',
            fontsize=13, fontweight='bold', color=C_DARK)
    services = [
        (0.3, 3.5, 'nginx\n:80/:443'),
        (2.5, 3.5, 'mac\n(FastAPI)\n:8000'),
        (4.7, 3.5, 'postgres\n:5432'),
        (6.9, 3.5, 'redis\n:6379'),
        (0.3, 1.5, 'qdrant\n:6333'),
        (2.5, 1.5, 'searxng\n:8080'),
        (4.7, 1.5, 'litellm\n:4000'),
        (6.9, 1.5, 'vLLM/Ollama\n:11434'),
    ]
    colors = [C_ORANGE, C_PRIMARY, C_SEC, C_AMBER, C_ACCENT, C_TEAL, C_GREEN, C_DARK]
    bw, bh = 2.0, 1.3
    for (x, y, label), color in zip(services, colors):
        _box(ax, x, y, bw, bh, label, color, 'white', 8)

    # Network
    _box(ax, 9.2, 2.5, 1.6, 2.3, 'mac-net\nbridge\nnetwork', '#e0e0e0', C_DARK, 7)

    # Arrows nginx→mac
    _arrow(ax, 2.3, 4.15, 2.5, 4.15, color=C_GRAY, lw=1.5)
    # mac→postgres, mac→redis
    _arrow(ax, 4.5, 4.15, 4.7, 4.15, color=C_GRAY, lw=1.2)
    _arrow(ax, 4.5, 3.7, 6.9, 3.7, color=C_GRAY, lw=1.2)
    # mac→qdrant, mac→searxng, mac→litellm
    _arrow(ax, 3.5, 3.5, 1.3, 2.8, color=C_GRAY, lw=1)
    _arrow(ax, 3.5, 3.5, 3.5, 2.8, color=C_GRAY, lw=1)
    _arrow(ax, 3.5, 3.5, 5.7, 2.8, color=C_GRAY, lw=1)
    # litellm→vllm
    _arrow(ax, 6.7, 2.15, 6.9, 2.15, color=C_GRAY, lw=1.2)

    return _to_buf(fig)


def dia_frontend_flow():
    fig, ax = _fig(12, 5, (0, 12), (0, 5))
    ax.text(6, 4.7, 'Frontend — User Journey Flow', ha='center', fontsize=13,
            fontweight='bold', color=C_DARK)

    pages = [
        (0.2, 2.5, 'Auth Page\nRoll No + DOB', C_TEAL),
        (2.3, 2.5, 'Set Password\n(first login)', C_AMBER),
        (4.4, 2.5, 'Dashboard\nStats · Models', C_PRIMARY),
        (6.5, 2.5, 'Chat\nStreaming AI', C_SEC),
        (8.6, 2.5, 'Settings\nProfile · Key', C_ACCENT),
        (10.0, 1.0, 'Admin Panel\nUsers · Registry', C_ORANGE),
    ]
    bw, bh = 1.9, 1.2
    for x, y, label, color in pages:
        _box(ax, x, y, bw, bh, label, color, 'white', 8)
    # Auth → Set Password
    _arrow(ax, 0.2 + bw, 3.1, 2.3, 3.1)
    # Set Password → Dashboard
    _arrow(ax, 2.3 + bw, 3.1, 4.4, 3.1)
    # Dashboard → Chat
    _arrow(ax, 4.4 + bw, 3.1, 6.5, 3.1)
    # Dashboard → Settings
    _arrow(ax, 4.4 + bw, 2.8, 8.6, 2.8, lw=1)
    # Dashboard → Admin
    _arrow(ax, 5.4, 2.5, 10.9, 1.0 + bh, color=C_ORANGE, lw=1)
    ax.text(8.0, 1.8, 'admin only', fontsize=7, color=C_ORANGE, fontstyle='italic')

    return _to_buf(fig)


def dia_search_flow():
    fig, ax = _fig(12, 4.5, (0, 12), (0, 4.5))
    ax.text(6, 4.2, 'Grounded Search Pipeline', ha='center', fontsize=13,
            fontweight='bold', color=C_DARK)
    steps = [
        (0.1, 'User\nQuery', C_TEAL),
        (2.3, 'SearXNG\nMeta-Search', C_ORANGE),
        (4.5, 'Top Results\nSnippets', C_AMBER),
        (6.7, 'LLM + Context\nGenerate Answer', C_PRIMARY),
        (9.2, 'Grounded\nResponse\n+ Citations', C_GREEN),
    ]
    bw, bh = 2, 1
    y = 1.5
    for x, label, color in steps:
        _box(ax, x, y, bw, bh, label, color, 'white', 9)
    for i in range(len(steps) - 1):
        _arrow(ax, steps[i][0] + bw, y + bh/2, steps[i+1][0], y + bh/2)
    _box(ax, 4.5, 0.2, 2, 0.6, 'Result Cache\n(Redis)', '#e8e8e8', C_DARK, 7, 'normal')
    _arrow(ax, 5.5, 1.5, 5.5, 0.85, color=C_GRAY, lw=1, style='<->')
    return _to_buf(fig)


def dia_test_coverage():
    fig, ax = plt.subplots(figsize=(9, 4.5))
    modules = ['Auth', 'Query', 'Models', 'Explore', 'Usage', 'Keys',
               'Quota', 'Guards', 'RAG', 'Search', 'Integration']
    tests = [15, 12, 8, 6, 8, 6, 5, 7, 8, 5, 6]
    colors_list = [C_PRIMARY, C_SEC, C_ACCENT, C_TEAL, C_GREEN,
                   C_AMBER, C_ORANGE, C_RED, C_DARK, C_SEC, C_ACCENT]
    bars = ax.barh(modules, tests, color=colors_list, height=0.55, edgecolor='white', linewidth=0.5)
    for bar, v in zip(bars, tests):
        ax.text(bar.get_width() + 0.2, bar.get_y() + bar.get_height()/2,
                f'{v}', va='center', fontsize=10, fontweight='bold', color=C_DARK)
    ax.set_xlabel('Number of Test Cases', fontsize=10, color=C_DARK)
    ax.set_title('Test Coverage Across Modules (81+ tests)', fontsize=13,
                 fontweight='bold', color=C_DARK, pad=12)
    ax.set_xlim(0, 20)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.tick_params(left=False, labelsize=9)
    ax.grid(axis='x', alpha=0.2)
    fig.tight_layout()
    return _to_buf(fig)


# ════════════════════════════════════════════════════════════════
#  MAIN DOCUMENT GENERATOR
# ════════════════════════════════════════════════════════════════

def generate_report():
    doc = new_doc()

    # ── TITLE PAGE ─────────────────────────────────────────
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('MAC')
    r.font.size = Pt(52)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1b, 0x49, 0x65)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('MBM AI Cloud')
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x2a, 0x6f, 0x97)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Comprehensive Project Report')
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor(0x0d, 0x1b, 0x2a)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Self-Hosted AI Inference Platform for MBM Engineering College')
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)

    for _ in range(4):
        doc.add_paragraph()

    info_lines = [
        ('Developed by:', 'Prof. Abhishek Gaur'),
        ('Department:', 'Computer Science & Engineering'),
        ('Institution:', 'MBM University, Jodhpur'),
        ('Date:', '08 April 2026'),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + '  ')
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)
        r2.bold = True
        r2.font.color.rgb = RGBColor(0x0d, 0x1b, 0x2a)

    doc.add_page_break()

    # ── TABLE OF CONTENTS (manual) ────────────────────────
    heading(doc, 'Table of Contents', level=1)

    toc_items = [
        ('1.', 'Executive Summary'),
        ('2.', 'Introduction & Motivation'),
        ('3.', 'System Architecture'),
        ('4.', 'Technology Stack'),
        ('5.', 'Module Breakdown'),
        ('  5.1', 'Authentication Module'),
        ('  5.2', 'Explore & Discovery Module'),
        ('  5.3', 'Query & Inference Module'),
        ('  5.4', 'Usage Tracking Module'),
        ('  5.5', 'Model Management Module'),
        ('  5.6', 'API Key Management Module'),
        ('  5.7', 'Quota & Rate Limiting Module'),
        ('  5.8', 'Guardrails Module'),
        ('  5.9', 'RAG (Retrieval-Augmented Generation) Module'),
        ('  5.10', 'Search & Grounded Answers Module'),
        ('  5.11', 'Integration & Worker Management Module'),
        ('6.', 'Database Design'),
        ('7.', 'Frontend — Progressive Web App'),
        ('8.', 'Security Architecture'),
        ('9.', 'Deployment & Infrastructure'),
        ('10.', 'Testing Strategy'),
        ('11.', 'API Endpoint Reference'),
        ('12.', 'Project Structure'),
        ('13.', 'Build Roadmap & Phase Summary'),
        ('14.', 'Future Scope'),
        ('15.', 'Conclusion'),
    ]
    for num, title_text in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(num + '  ')
        r1.font.size = Pt(11)
        r1.bold = True
        r2 = p.add_run(title_text)
        r2.font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════
    heading(doc, '1. Executive Summary', level=1)

    body(doc, 'MAC (MBM AI Cloud) is a fully self-hosted, zero-cloud AI inference platform purpose-built for '
         'MBM Engineering College, Jodhpur. The platform provides students and faculty with on-premise access to '
         'state-of-the-art large language models — for code generation, mathematical reasoning, general text tasks, '
         'image understanding, and speech transcription — through a clean, well-documented REST API and a modern '
         'Progressive Web App (PWA) frontend.')

    body(doc, 'The system is designed around a core principle: no student data ever leaves the college network, '
         'and no cloud API subscription is required. All inference runs on college-owned GPU hardware. The platform '
         'starts on a single PC and scales horizontally to 30+ nodes with zero code changes — only new vLLM worker '
         'containers are registered.')

    heading(doc, 'Key Highlights', level=2)
    highlights = [
        ('50+ REST API endpoints', ' spanning 11 modules — auth, explore, query, usage, models, keys, quota, guardrails, RAG, search, and integration.'),
        ('OpenAI-compatible API', ' — students use the official OpenAI Python SDK with only a base_url change.'),
        ('Unified DOB-based authentication', ' — students verify using their Registration Number and Date of Birth from the college registry; no separate signup needed.'),
        ('Content guardrails', ' — input checks for prompt injection, harmful content, academic dishonesty; output checks for PII redaction.'),
        ('RAG pipeline', ' — document ingestion, chunking, vector embedding via Qdrant, and retrieval-augmented generation.'),
        ('Grounded web search', ' — SearXNG meta-search integrated with LLM for cited, factual answers.'),
        ('PWA frontend', ' — installable app with dashboard, multi-session chat (SSE streaming), settings, and admin panel.'),
        ('81+ automated tests', ' across all 11 modules with pytest-asyncio — 100% passing.'),
        ('Docker Compose', ' one-command deployment with 7 containerised services.'),
    ]
    for bold_part, rest in highlights:
        bullet(doc, rest, bold_prefix=bold_part)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 2. INTRODUCTION & MOTIVATION
    # ══════════════════════════════════════════════════════════
    heading(doc, '2. Introduction & Motivation', level=1)

    body(doc, 'The rapid advancement of Large Language Models (LLMs) has created a fundamental shift in how software is '
         'developed, studied, and taught. Students at engineering colleges increasingly rely on AI tools for code generation, '
         'debugging, mathematical reasoning, and research assistance. However, commercial AI services such as OpenAI GPT-4, '
         'Google Gemini, and Anthropic Claude impose significant per-token costs that are prohibitive for educational institutions '
         'serving hundreds of students.')

    body(doc, 'MBM Engineering College possesses substantial computing resources in its laboratories — specifically, '
         'NVIDIA RTX 3060 (12 GB VRAM) GPUs across multiple lab PCs. These machines sit idle after lab hours and during '
         'weekends. MAC repurposes this existing hardware to serve state-of-the-art open-source AI models, eliminating '
         'the need for any cloud subscriptions.')

    heading(doc, 'Problem Statement', level=2)
    body(doc, 'Design and implement a self-hosted AI inference platform that provides MBM students and faculty with '
         'free, private, and high-performance access to multiple AI models through a standardised API interface, '
         'with built-in authentication, usage control, content safety guardrails, and knowledge retrieval capabilities.')

    heading(doc, 'Objectives', level=2)
    objectives = [
        'Provide zero-cost AI model access to all students on the college LAN.',
        'Ensure complete data privacy — all queries and responses stay within the college network.',
        'Support multiple model types: text generation, code assistance, mathematical reasoning, image understanding, and speech transcription.',
        'Expose an OpenAI-compatible API so existing SDKs and tutorials work without modification.',
        'Implement robust authentication, rate limiting, and content safety guardrails.',
        'Build a Progressive Web App for students who prefer a ChatGPT-like interface.',
        'Enable Retrieval-Augmented Generation (RAG) over college textbooks and lecture notes.',
        'Design for horizontal scalability — start on one PC, expand to N PCs with no code changes.',
    ]
    for obj in objectives:
        bullet(doc, obj)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 3. SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════
    heading(doc, '3. System Architecture', level=1)

    body(doc, 'The MAC platform follows a layered microservices architecture. Requests flow from the client through '
         'a reverse proxy, into the API gateway (which handles authentication, rate limiting, and routing), then to '
         'the appropriate service — either the LLM inference engine, the vector database for RAG, the web search engine, '
         'or the relational database for user/usage data.')

    add_diagram(doc, dia_architecture())

    body(doc, 'Figure 1: Complete system architecture showing the request flow from the student client through Nginx, '
         'the FastAPI gateway, and into the backend services.', italic=True, size=10, justify=False)

    heading(doc, 'Request Flow', level=2)
    body(doc, 'A typical chat request follows this path:')
    steps = [
        'The student opens the PWA or sends an API request from their code.',
        'Nginx receives the request, terminates TLS, and forwards it to the FastAPI gateway.',
        'The FastAPI middleware validates the JWT token or API key.',
        'The rate limiter checks hourly request count and daily token quota.',
        'Input guardrails scan the prompt for injection attempts and harmful content.',
        'The query router forwards the request to LiteLLM, which routes to the appropriate vLLM worker.',
        'vLLM performs GPU-accelerated inference and returns the response.',
        'Output guardrails scan the response for PII and harmful content.',
        'Usage is logged to the database and the response is returned to the student.',
    ]
    for i, step in enumerate(steps, 1):
        bullet(doc, f'{step}')

    add_diagram(doc, dia_request_pipeline())
    body(doc, 'Figure 2: Request processing pipeline showing validation stages and rejection points.', italic=True, size=10, justify=False)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 4. TECHNOLOGY STACK
    # ══════════════════════════════════════════════════════════
    heading(doc, '4. Technology Stack', level=1)

    body(doc, 'Each technology was evaluated against alternatives before selection. The guiding criteria were: '
         'open-source licensing, self-hostability, async performance, and community support.')

    styled_table(doc, ['Layer', 'Technology', 'Version', 'Rationale'],
    [
        ['API Framework', 'FastAPI', '0.115.6', 'Async, auto-generated OpenAPI docs, type-safe with Pydantic'],
        ['Runtime', 'Python', '3.11+', 'Rich AI/ML ecosystem, asyncio support, broad library availability'],
        ['ASGI Server', 'Uvicorn', '0.34.0', 'High-performance, supports HTTP/1.1 and WebSockets'],
        ['Database (Prod)', 'PostgreSQL 16', '-', 'ACID-compliant, excellent JSON support, concurrent writes'],
        ['Database (Dev)', 'SQLite + aiosqlite', '0.20.0', 'Zero-config for development, file-based, async support'],
        ['ORM', 'SQLAlchemy', '2.0.36', 'Async engine, mature, supports PostgreSQL and SQLite'],
        ['Cache', 'Redis 7', '-', 'In-memory, sub-ms latency, rate limiting primitives'],
        ['LLM Inference', 'vLLM / Ollama', '-', 'PagedAttention, continuous batching, OpenAI-compatible API'],
        ['Model Router', 'LiteLLM', '-', 'Unified proxy, load balancing, health checks, fallback chains'],
        ['Vector DB', 'Qdrant', '-', 'Purpose-built for embeddings, fast similarity search'],
        ['Web Search', 'SearXNG', '-', 'Self-hosted meta-search, no API keys required'],
        ['Reverse Proxy', 'Nginx', '-', 'TLS termination, SSE streaming, request buffering'],
        ['Auth Tokens', 'python-jose (JWT)', '-', 'HS256 signing, stateless authentication'],
        ['Password Hashing', 'bcrypt', '4.2.1', 'Industry standard, async-wrapped to avoid blocking'],
        ['HTTP Client', 'httpx', '0.28.1', 'Async, streaming support, connection pooling'],
        ['SSE', 'sse-starlette', '2.2.1', 'Server-Sent Events for streaming LLM responses'],
        ['Containers', 'Docker + Compose', '-', 'Reproducible deployment, one-command startup'],
        ['Frontend', 'Vanilla JS PWA', '-', 'Zero build step, installable, offline-capable'],
        ['Charts', 'Chart.js', '4.4.7', 'Lightweight, responsive donut charts for dashboard'],
        ['Testing', 'pytest + pytest-asyncio', '-', 'Async test support, fixtures, parameterisation'],
    ], col_widths=[1.3, 1.5, 0.6, 3.2])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 5. MODULE BREAKDOWN
    # ══════════════════════════════════════════════════════════
    heading(doc, '5. Module Breakdown', level=1)

    body(doc, 'The MAC codebase is organised into 11 router modules, 7 service modules, 2 middleware modules, '
         '4 ORM model files, and 11 Pydantic schema files. Each module has a clear single responsibility.')

    add_diagram(doc, dia_module_structure())
    body(doc, 'Figure 3: Complete module architecture showing routers, services, middleware, models, and schemas.', italic=True, size=10, justify=False)

    # ── 5.1 Authentication ──
    heading(doc, '5.1 Authentication Module', level=2)
    body(doc, 'The authentication module is the most complex, handling the complete user lifecycle from first-time '
         'DOB verification through to admin user management.')

    heading(doc, 'Unified Auth Flow', level=3)
    body(doc, 'Unlike traditional systems that separate login and signup, MAC uses a single unified flow. Students '
         'enter their Registration Number and Date of Birth (in DDMMYYYY format). The system verifies these against '
         'the pre-loaded Student Registry in the database. If the student has not logged in before, an account is '
         'automatically created with must_change_password set to true, redirecting them to set a password. On subsequent '
         'visits, the same flow authenticates them.')

    add_diagram(doc, dia_auth_flow())
    body(doc, 'Figure 4: Unified authentication flow — from DOB verification to token issuance.', italic=True, size=10, justify=False)

    heading(doc, 'Security Features', level=3)
    security_features = [
        ('Password hashing:', ' bcrypt with work factor 12. All bcrypt calls are wrapped with asyncio.to_thread() to prevent blocking the async event loop.'),
        ('JWT tokens:', ' HS256-signed access tokens with 24-hour expiry. Refresh tokens are random 32-byte secrets stored as SHA-256 hashes.'),
        ('Account lockout:', ' After 5 failed login attempts, the account is locked for 15 minutes.'),
        ('API keys:', ' Each user receives a unique api_key (mac_sk_live_*) for programmatic access. Keys are stored in the database.'),
        ('Role-based access:', ' Three roles — student, faculty, admin — with escalating privileges.'),
    ]
    for bold_part, rest in security_features:
        bullet(doc, rest, bold_prefix=bold_part)

    heading(doc, 'Admin Endpoints', level=3)
    body(doc, 'Administrators can manage users, the student registry, and view platform statistics through 10 dedicated '
         'admin endpoints under /auth/admin/*:')
    styled_table(doc, ['Endpoint', 'Description'],
    [
        ['GET /auth/admin/users', 'List all registered users with roles and status'],
        ['POST /auth/admin/users', 'Create a new user account'],
        ['PUT /auth/admin/users/{id}/role', 'Change user role (student/faculty/admin)'],
        ['PUT /auth/admin/users/{id}/status', 'Activate or deactivate a user account'],
        ['DELETE /auth/admin/users/{id}', 'Permanently delete a user'],
        ['POST /auth/admin/users/{id}/reset-password', 'Reset user password to a temporary value'],
        ['POST /auth/admin/users/{id}/regenerate-key', 'Regenerate a user\'s API key'],
        ['GET /auth/admin/registry', 'List all entries in the student registry'],
        ['POST /auth/admin/registry', 'Add a new student to the registry'],
        ['POST /auth/admin/registry/bulk', 'Bulk import students from CSV/JSON data'],
        ['GET /auth/admin/stats', 'Dashboard statistics (users, requests, tokens)'],
    ], col_widths=[3.0, 3.6])

    # ── 5.2 Explore ──
    heading(doc, '5.2 Explore & Discovery Module', level=2)
    body(doc, 'The explore module provides read-only discovery endpoints that require no authentication. Students '
         'can see what models are available, what capabilities the platform offers, search models by tag, list all '
         'API endpoints, and check platform health — all before writing any code.')

    styled_table(doc, ['Endpoint', 'Description'],
    [
        ['GET /explore/models', 'List all deployed models with status, capabilities, and context length'],
        ['GET /explore/models/search?tag=code', 'Filter models by capability tag (code, math, vision, chat)'],
        ['GET /explore/models/{model_id}', 'Detailed model info including benchmarks and example prompts'],
        ['GET /explore/endpoints', 'Auto-generated list of all 50+ API endpoints in the platform'],
        ['GET /explore/health', 'Platform health: uptime, models loaded, version number'],
    ], col_widths=[3.2, 3.4])

    # ── 5.3 Query ──
    heading(doc, '5.3 Query & Inference Module', level=2)
    body(doc, 'The query module is the core inference API. It supports chat completions, text completions, embeddings, '
         're-ranking, vision analysis, and speech-to-text. All endpoints are OpenAI-compatible — students can use the '
         'official OpenAI SDK by simply changing the base_url.')

    heading(doc, 'Smart Routing', level=3)
    body(doc, 'When the model parameter is set to "auto", the system analyses the request content and routes to the '
         'optimal model:')
    styled_table(doc, ['Content Signal', 'Routed To'],
    [
        ['Code keywords (function, debug, class, compile, etc.)', 'qwen2.5-coder:7b'],
        ['Math/reasoning (solve, prove, equation, step-by-step)', 'deepseek-r1:8b'],
        ['Image attachment in request', 'llava:7b'],
        ['General text, Q&A, summarisation', 'qwen2.5:14b'],
    ])

    heading(doc, 'Streaming', level=3)
    body(doc, 'When stream=true is set in the request, the response is delivered as Server-Sent Events (SSE). Each '
         'event contains a delta with a partial token, enabling real-time display in the frontend. The Nginx '
         'configuration has proxy_buffering disabled to support SSE passthrough.')

    heading(doc, 'Endpoints', level=3)
    styled_table(doc, ['Endpoint', 'Content-Type', 'Description'],
    [
        ['POST /query/chat', 'application/json', 'Chat completion — multi-turn, streaming/non-streaming'],
        ['POST /query/completions', 'application/json', 'Raw text completion'],
        ['POST /query/embeddings', 'application/json', 'Text → 768-dim vector embeddings'],
        ['POST /query/rerank', 'application/json', 'Re-rank passages by cosine similarity'],
        ['POST /query/vision', 'multipart/form-data', 'Image + text → answer'],
        ['POST /query/speech-to-text', 'multipart/form-data', 'Audio → transcribed text'],
    ], col_widths=[2.2, 1.8, 2.6])

    # ── 5.4 Usage ──
    heading(doc, '5.4 Usage Tracking Module', level=2)
    body(doc, 'Every API call is logged with the model used, token counts (input and output separately), latency in '
         'milliseconds, status code, and a unique request ID. Students can track their own consumption, and administrators '
         'can monitor the entire platform.')

    styled_table(doc, ['Endpoint', 'Auth', 'Description'],
    [
        ['GET /usage/me', 'Any user', 'Token usage — today, this week, this month, by model'],
        ['GET /usage/me/history', 'Any user', 'Paginated request history with timestamps'],
        ['GET /usage/me/quota', 'Any user', 'Current quota limits and remaining balance'],
        ['GET /usage/admin/all', 'Admin', 'All users\' usage summary'],
        ['GET /usage/admin/user/{roll}', 'Admin', 'Specific student\'s detailed usage'],
        ['GET /usage/admin/models', 'Admin', 'Per-model usage stats across the platform'],
    ], col_widths=[2.5, 0.8, 3.3])

    # ── 5.5 Models ──
    heading(doc, '5.5 Model Management Module', level=2)
    body(doc, 'This module provides model lifecycle management — listing available models, loading/unloading models '
         'from GPU memory, checking model health, and downloading new models from Ollama.')

    styled_table(doc, ['Model', 'Parameters', 'Specialty', 'VRAM'],
    [
        ['Qwen2.5-Coder 7B', '7B', 'Code generation, debugging, completion', '~5 GB'],
        ['DeepSeek-R1 8B', '8B', 'Mathematics, step-by-step reasoning', '~6 GB'],
        ['LLaVA 1.6 7B', '7B', 'Image understanding, visual Q&A', '~8 GB'],
        ['Qwen2.5 14B', '14B', 'General text, summarisation, Q&A', '~10 GB'],
        ['Whisper Large v3', '1.5B', 'Speech-to-text transcription', '~3 GB'],
    ])

    # ── 5.6 Keys ──
    heading(doc, '5.6 API Key Management Module', level=2)
    body(doc, 'Every user receives a unique API key (format: mac_sk_live_<hex>) upon account creation. The module '
         'handles key generation, revocation, usage statistics, and admin-level key management. API keys serve as an '
         'alternative authentication method to JWT tokens, particularly useful for programmatic access from scripts and notebooks.')

    styled_table(doc, ['Endpoint', 'Description'],
    [
        ['GET /keys/my-key', 'Get current API key (partially masked) and metadata'],
        ['POST /keys/generate', 'Generate new API key (invalidates old one)'],
        ['GET /keys/my-key/stats', 'Token consumption against this key'],
        ['DELETE /keys/my-key', 'Revoke current key permanently'],
        ['GET /keys/admin/all', 'List all student API keys and status (admin)'],
        ['POST /keys/admin/revoke', 'Force-revoke a student\'s key (admin)'],
    ], col_widths=[2.5, 4.1])

    # ── 5.7 Quota ──
    heading(doc, '5.7 Quota & Rate Limiting Module', level=2)
    body(doc, 'Rate limits prevent any single user from monopolising GPU resources. The system applies a sliding-window '
         'algorithm for request rate limiting and a daily token counter for quota enforcement. Administrators can set '
         'per-user quota overrides.')

    styled_table(doc, ['Role', 'Daily Token Limit', 'Requests/Hour', 'Max Tokens/Request'],
    [
        ['Student', '50,000', '100', '4,096'],
        ['Faculty', '200,000', '500', '8,192'],
        ['Admin', 'Unlimited', 'Unlimited', '16,384'],
    ])

    body(doc, 'Rate limit information is returned in every API response via custom headers: X-RateLimit-Limit, '
         'X-RateLimit-Remaining, X-RateLimit-Reset, X-TokenLimit-Limit, and X-TokenLimit-Remaining.')

    # ── 5.8 Guardrails ──
    heading(doc, '5.8 Guardrails Module', level=2)
    body(doc, 'The guardrails module provides content filtering for both inputs and outputs. It is the safety layer '
         'that protects against prompt injection, harmful content generation, and inadvertent PII exposure.')

    add_diagram(doc, dia_guardrail_pipeline())
    body(doc, 'Figure 5: Input and output guardrail pipeline with check categories.', italic=True, size=10, justify=False)

    heading(doc, 'Input Checks', level=3)
    styled_table(doc, ['Category', 'What It Detects', 'Action'],
    [
        ['Prompt Injection', 'Override attempts (ignore instructions, you are now, etc.)', 'Block'],
        ['Harmful Content', 'Violence, weapons, drugs, explicit content keywords', 'Block'],
        ['Academic Dishonesty', 'Exact exam answers, plagiarism requests', 'Flag + warn'],
        ['Max Length', 'Prompts exceeding 32,000 characters', 'Truncate'],
    ])

    heading(doc, 'Output Checks (PII Redaction)', level=3)
    styled_table(doc, ['Pattern', 'Example', 'Action'],
    [
        ['Email addresses', 'user@example.com → [EMAIL REDACTED]', 'Redact'],
        ['Indian phone numbers', '+91-98765-43210 → [PHONE REDACTED]', 'Redact'],
        ['ID/card numbers', '12-digit numbers → [ID REDACTED]', 'Redact'],
    ])

    body(doc, 'Built-in patterns are always active. Administrators can add custom regex patterns via the '
         'PUT /guardrails/rules endpoint, which are stored in the database and applied alongside the defaults.')

    # ── 5.9 RAG ──
    heading(doc, '5.9 RAG (Retrieval-Augmented Generation) Module', level=2)
    body(doc, 'The RAG module enables students and faculty to upload documents (textbooks, lecture notes, research papers) '
         'and query them using natural language. The system chunks documents, generates vector embeddings, stores them in '
         'Qdrant, and retrieves the most relevant passages when answering questions.')

    add_diagram(doc, dia_rag_pipeline())
    body(doc, 'Figure 6: RAG pipeline showing document ingestion (top) and query-time retrieval (bottom).', italic=True, size=10, justify=False)

    heading(doc, 'Document Ingestion', level=3)
    steps_rag = [
        'User uploads a document via POST /rag/ingest (multipart/form-data).',
        'The text is extracted and split into chunks of ~512 tokens with 50-token overlap at word boundaries.',
        'Each chunk is passed through the embedding model to generate a 768-dimensional vector.',
        'Vectors are stored in Qdrant under the specified collection with metadata (document ID, chunk index).',
        'The document record is saved in PostgreSQL with chunk_count and status.',
    ]
    for step in steps_rag:
        bullet(doc, step)

    heading(doc, 'Query Flow', level=3)
    body(doc, 'When a student queries via POST /rag/query, the question is embedded, Qdrant performs a similarity '
         'search to find the top-k most relevant chunks, and these chunks are injected into the LLM prompt as context. '
         'The LLM generates an answer grounded in the retrieved content, with source citations included.')

    # ── 5.10 Search ──
    heading(doc, '5.10 Search & Grounded Answers Module', level=2)
    body(doc, 'This module integrates a self-hosted SearXNG meta-search engine with the LLM to provide web-grounded, '
         'cited answers. Unlike RAG (which searches local documents), this module searches the open web.')

    add_diagram(doc, dia_search_flow())
    body(doc, 'Figure 7: Grounded search pipeline — from query to cited answer.', italic=True, size=10, justify=False)

    styled_table(doc, ['Endpoint', 'Description'],
    [
        ['POST /search/web', 'Web search via SearXNG (returns top results with snippets)'],
        ['POST /search/wikipedia', 'Wikipedia article search'],
        ['POST /search/grounded', 'Web search + LLM-generated answer with inline citations'],
        ['GET /search/cache', 'List cached search results (1-hour in-memory cache)'],
    ], col_widths=[2.5, 4.1])

    # ── 5.11 Integration ──
    heading(doc, '5.11 Integration & Worker Management Module', level=2)
    body(doc, 'The integration module manages the connection between the FastAPI gateway and the LLM inference workers. '
         'It exposes smart routing rules, worker node status, and the inference queue.')

    styled_table(doc, ['Endpoint', 'Description'],
    [
        ['GET /integration/routing-rules', 'View current smart routing configuration'],
        ['PUT /integration/routing-rules', 'Update routing rules (admin)'],
        ['GET /integration/workers', 'List all worker nodes with status'],
        ['GET /integration/workers/{node_id}', 'Detailed worker info (GPU temp, VRAM, models loaded)'],
        ['POST /integration/workers/{node_id}/drain', 'Drain a worker for maintenance (admin)'],
        ['GET /integration/queue', 'Current inference queue depth and wait times'],
    ], col_widths=[3.2, 3.4])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 6. DATABASE DESIGN
    # ══════════════════════════════════════════════════════════
    heading(doc, '6. Database Design', level=1)

    body(doc, 'The relational schema is implemented using SQLAlchemy 2.0 with async support. In development, SQLite '
         'is used (via aiosqlite) for zero-configuration setup; in production, PostgreSQL 16 is used (via asyncpg). '
         'All primary keys use UUIDs to prevent sequential ID enumeration attacks.')

    add_diagram(doc, dia_db_schema())
    body(doc, 'Figure 8: Entity-Relationship diagram showing all database tables and their relationships.', italic=True, size=10, justify=False)

    heading(doc, 'Table Descriptions', level=2)

    styled_table(doc, ['Table', 'Records', 'Key Columns', 'Purpose'],
    [
        ['student_registry', 'Pre-loaded', 'roll_number, name, dob, department, batch_year', 'College student records for DOB verification'],
        ['users', 'Dynamic', 'roll_number, email, role, password_hash, api_key, must_change_password', 'Registered platform users with auth credentials'],
        ['refresh_tokens', 'Dynamic', 'user_id, token_hash, expires_at, revoked', 'JWT refresh tokens for session continuity'],
        ['usage_logs', 'Append-only', 'user_id, model, tokens_in, tokens_out, latency_ms', 'Per-request consumption tracking'],
        ['quota_overrides', 'Admin-set', 'user_id, daily_tokens, requests_per_hour', 'Per-user quota overrides beyond role defaults'],
        ['guardrail_rules', 'Admin-set', 'category, action, pattern, enabled, priority', 'Custom content filter rules'],
        ['rag_collections', 'Dynamic', 'name, description, document_count', 'Named document collections for RAG'],
        ['rag_documents', 'Dynamic', 'collection_id, title, chunk_count, status', 'Uploaded documents with processing status'],
    ], col_widths=[1.3, 0.8, 2.8, 1.7])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 7. FRONTEND — PWA
    # ══════════════════════════════════════════════════════════
    heading(doc, '7. Frontend — Progressive Web App', level=1)

    body(doc, 'The MAC frontend is a single-page Progressive Web App (PWA) built with vanilla JavaScript, HTML, and CSS — '
         'no build step, no Node.js, no React or Angular dependency. The entire frontend fits in four files: index.html, '
         'app.js (~750 lines), style.css, and manifest.json. It is served directly by the FastAPI static file mount.')

    heading(doc, 'Key Features', level=2)
    frontend_features = [
        ('Installable PWA:', ' Users can install the app to their desktop or phone via the browser\'s "Install" prompt. A service worker (sw.js) enables offline caching.'),
        ('Single-page routing:', ' Navigation is handled via hash-based routing (#dashboard, #chat, #settings, #admin) with no page reloads.'),
        ('Unified auth page:', ' A single form for Registration Number + Date of Birth replaces traditional login/signup flows.'),
        ('Multi-session chat:', ' Students can create multiple chat sessions, each stored in localStorage. Messages stream in real-time via SSE.'),
        ('Dashboard:', ' Displays welcome card, today\'s tokens, requests per hour, chat session count, two donut charts (token usage and request breakdown), model status list, and recent activity table.'),
        ('Admin panel:', ' Three-tab interface — Overview (stats), Users (full CRUD table), Student Registry (list/add/bulk import).'),
        ('Settings:', ' Profile editing (name, email), password change with eye-toggle visibility, API key display and copy.'),
        ('XSS protection:', ' All user-generated content is HTML-escaped via the esc() utility function before DOM injection.'),
    ]
    for bold_part, rest in frontend_features:
        bullet(doc, rest, bold_prefix=bold_part)

    add_diagram(doc, dia_frontend_flow())
    body(doc, 'Figure 9: Frontend user journey — from authentication to dashboard, chat, settings, and admin.', italic=True, size=10, justify=False)

    heading(doc, 'Design', level=2)
    body(doc, 'The frontend uses a black-and-white professional theme with CSS custom properties for colours. '
         'The "MAC" branding features a glitch animation effect. The layout is fully responsive with breakpoints '
         'at 768px and 600px for tablet and mobile devices.')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 8. SECURITY ARCHITECTURE
    # ══════════════════════════════════════════════════════════
    heading(doc, '8. Security Architecture', level=1)

    body(doc, 'Security is implemented at every layer of the stack, following the OWASP Top 10 guidelines:')

    styled_table(doc, ['Threat / Concern', 'Mitigation'],
    [
        ['Password storage', 'bcrypt (work factor 12), async-wrapped to prevent event loop blocking'],
        ['JWT token theft', 'Short expiry (24h), refresh tokens stored as SHA-256 hashes, revocable'],
        ['API key exposure', 'Keys shown only once at creation; masked in all subsequent reads'],
        ['Brute force login', 'Account lockout after 5 failed attempts for 15 minutes'],
        ['SQL injection', 'SQLAlchemy ORM with parameterised queries throughout'],
        ['XSS (frontend)', 'esc() function HTML-escapes all dynamic content before DOM insertion'],
        ['CSRF', 'API uses stateless JWT; no cookies for auth. CORS restricted to frontend origin'],
        ['Prompt injection', 'Input guardrails detect and block override/jailbreak patterns'],
        ['PII leakage', 'Output guardrails redact emails, phone numbers, and ID numbers'],
        ['Rate abuse', 'Sliding-window rate limiter per user + daily token quota enforcement'],
        ['Transport security', 'Nginx terminates TLS; internal traffic on Docker bridge network'],
        ['Role escalation', 'require_admin dependency on all admin endpoints; role checked from JWT'],
        ['Excessive payload', 'Nginx enforces 25 MB upload limit; Pydantic validates all request schemas'],
    ], col_widths=[1.8, 4.8])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 9. DEPLOYMENT & INFRASTRUCTURE
    # ══════════════════════════════════════════════════════════
    heading(doc, '9. Deployment & Infrastructure', level=1)

    body(doc, 'The entire platform is containerised with Docker and orchestrated with Docker Compose. A single '
         'docker-compose up command starts all 7 services on any machine with Docker installed.')

    add_diagram(doc, dia_deployment())
    body(doc, 'Figure 10: Docker Compose deployment showing all 7 services and the bridge network.', italic=True, size=10, justify=False)

    heading(doc, 'Service Inventory', level=2)
    styled_table(doc, ['Service', 'Image', 'Port', 'Purpose'],
    [
        ['mac', 'Custom (Dockerfile)', '8000', 'FastAPI API gateway'],
        ['postgres', 'postgres:16-alpine', '5432', 'Relational database'],
        ['redis', 'redis:7-alpine', '6379', 'Cache and rate limiting'],
        ['nginx', 'nginx:alpine', '80/443', 'Reverse proxy and TLS'],
        ['qdrant', 'qdrant/qdrant', '6333', 'Vector database for RAG'],
        ['searxng', 'searxng/searxng', '8080', 'Self-hosted web search'],
        ['litellm', 'ghcr.io/berriai/litellm', '4000', 'LLM proxy and model router'],
    ])

    heading(doc, 'Scaling Strategy', level=2)
    body(doc, 'To scale from 1 PC to N PCs:')
    scaling_steps = [
        'Install vLLM or Ollama on additional lab PCs with GPUs.',
        'Load the desired models on each new node.',
        'Register the new nodes in litellm/config.yaml with their IP addresses.',
        'Restart the LiteLLM proxy (or hot-reload the config).',
        'LiteLLM automatically load-balances requests across all registered nodes.',
    ]
    for step in scaling_steps:
        bullet(doc, step)

    body(doc, 'The FastAPI gateway, PostgreSQL database, Redis cache, and Nginx proxy remain on the primary node. '
         'Only the GPU inference workers are distributed. This architecture supports 30+ worker nodes with a single '
         'gateway instance.')

    heading(doc, 'Development Mode', level=2)
    body(doc, 'For development, the platform can run without Docker:')
    code_block(doc,
        '# Create virtual environment and install dependencies\n'
        'python -m venv venv\n'
        '.\\venv\\Scripts\\Activate.ps1    # Windows\n'
        'pip install -r requirements.txt\n\n'
        '# Run the server (SQLite, no PostgreSQL/Redis needed)\n'
        'set PYTHONPATH=D:\\MBM\n'
        'uvicorn mac.main:app --host 0.0.0.0 --port 8000 --reload'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 10. TESTING STRATEGY
    # ══════════════════════════════════════════════════════════
    heading(doc, '10. Testing Strategy', level=1)

    body(doc, 'The project includes a comprehensive automated test suite with 81+ test cases across all 11 modules. '
         'Tests are written using pytest with the pytest-asyncio plugin for testing async FastAPI endpoints. An async '
         'test client (httpx.AsyncClient) makes real HTTP requests against the in-memory ASGI app — no external server needed.')

    heading(doc, 'Test Infrastructure', level=2)
    test_infra = [
        ('Fixtures:', ' conftest.py provides auto-use database setup (create/drop tables per test), test_user, admin_user, and pre-authenticated header fixtures.'),
        ('Database isolation:', ' Each test runs against a fresh in-memory SQLite database to prevent cross-test contamination.'),
        ('No mocking:', ' Tests hit the actual FastAPI application through ASGI transport — the same code path as production.'),
    ]
    for bold_part, rest in test_infra:
        bullet(doc, rest, bold_prefix=bold_part)

    heading(doc, 'Test Coverage by Module', level=2)
    add_diagram(doc, dia_test_coverage())
    body(doc, 'Figure 11: Test case distribution across all 11 modules.', italic=True, size=10, justify=False)

    styled_table(doc, ['Test File', 'Module', 'Tests', 'What Is Tested'],
    [
        ['test_auth.py', 'Authentication', '15', 'Login, logout, refresh, DOB verify, admin CRUD, registry'],
        ['test_query.py', 'Query', '12', 'Chat, completions, embeddings, rerank, vision, streaming'],
        ['test_models.py', 'Models', '8', 'List, detail, health, load/unload, download'],
        ['test_explore.py', 'Explore', '6', 'Models, endpoints, health, search by tag'],
        ['test_usage.py', 'Usage', '8', 'My usage, history, quota, admin views'],
        ['test_keys.py', 'Keys', '6', 'Generate, view, revoke, admin list'],
        ['test_quota.py', 'Quota', '5', 'Limits, personal quota, admin overrides'],
        ['test_guardrails.py', 'Guardrails', '7', 'Input filtering, PII redaction, rule management'],
        ['test_rag.py', 'RAG', '8', 'Ingest, query, collections, document lifecycle'],
        ['test_search.py', 'Search', '5', 'Web search, Wikipedia, grounded search, cache'],
        ['test_integration.py', 'Integration', '6', 'Routing rules, workers, queue status'],
    ], col_widths=[1.5, 1.2, 0.5, 3.4])

    heading(doc, 'Running Tests', level=2)
    code_block(doc,
        '# Run all 81+ tests\n'
        'set PYTHONPATH=D:\\MBM\n'
        'pytest tests/ -v\n\n'
        '# Run a specific module\n'
        'pytest tests/test_auth.py -v\n\n'
        '# Run with coverage report\n'
        'pytest tests/ --cov=mac --cov-report=term-missing'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 11. API ENDPOINT REFERENCE
    # ══════════════════════════════════════════════════════════
    heading(doc, '11. API Endpoint Reference', level=1)

    body(doc, 'Complete listing of all REST API endpoints exposed by the MAC platform. Base URL: http://<server-ip>/api/v1')

    all_endpoints = [
        # Auth
        ('POST', '/auth/verify', 'No', 'Unified DOB-based authentication'),
        ('POST', '/auth/login', 'No', 'Password-based login (legacy/API)'),
        ('POST', '/auth/set-password', 'JWT', 'First-time password setup'),
        ('POST', '/auth/logout', 'JWT', 'Revoke session tokens'),
        ('POST', '/auth/refresh', 'No', 'Exchange refresh token'),
        ('GET', '/auth/me', 'JWT', 'User profile with quota info'),
        ('PUT', '/auth/me/profile', 'JWT', 'Update name/email'),
        ('POST', '/auth/change-password', 'JWT', 'Change password'),
        ('GET', '/auth/admin/users', 'Admin', 'List all users'),
        ('POST', '/auth/admin/users', 'Admin', 'Create user'),
        ('PUT', '/auth/admin/users/{id}/role', 'Admin', 'Change role'),
        ('PUT', '/auth/admin/users/{id}/status', 'Admin', 'Toggle active'),
        ('DELETE', '/auth/admin/users/{id}', 'Admin', 'Delete user'),
        ('POST', '/auth/admin/users/{id}/reset-password', 'Admin', 'Reset password'),
        ('POST', '/auth/admin/users/{id}/regenerate-key', 'Admin', 'Regenerate API key'),
        ('GET', '/auth/admin/registry', 'Admin', 'List registry'),
        ('POST', '/auth/admin/registry', 'Admin', 'Add registry entry'),
        ('POST', '/auth/admin/registry/bulk', 'Admin', 'Bulk import'),
        ('GET', '/auth/admin/stats', 'Admin', 'Dashboard stats'),
        # Explore
        ('GET', '/explore/models', 'No', 'List models'),
        ('GET', '/explore/models/search', 'No', 'Search by tag'),
        ('GET', '/explore/models/{id}', 'No', 'Model detail'),
        ('GET', '/explore/endpoints', 'No', 'List all endpoints'),
        ('GET', '/explore/health', 'No', 'Platform health'),
        # Query
        ('POST', '/query/chat', 'JWT/Key', 'Chat completion'),
        ('POST', '/query/completions', 'JWT/Key', 'Text completion'),
        ('POST', '/query/embeddings', 'JWT/Key', 'Vector embeddings'),
        ('POST', '/query/rerank', 'JWT/Key', 'Re-rank passages'),
        ('POST', '/query/vision', 'JWT/Key', 'Image analysis'),
        ('POST', '/query/speech-to-text', 'JWT/Key', 'Speech to text'),
        # Usage
        ('GET', '/usage/me', 'JWT', 'My usage stats'),
        ('GET', '/usage/me/history', 'JWT', 'Request history'),
        ('GET', '/usage/me/quota', 'JWT', 'My quota status'),
        ('GET', '/usage/admin/all', 'Admin', 'All users usage'),
        ('GET', '/usage/admin/user/{roll}', 'Admin', 'User usage detail'),
        ('GET', '/usage/admin/models', 'Admin', 'Per-model stats'),
        # Models
        ('GET', '/models/', 'No', 'List models'),
        ('GET', '/models/{id}', 'No', 'Model detail'),
        ('POST', '/models/{id}/load', 'Admin', 'Load model'),
        ('POST', '/models/{id}/unload', 'Admin', 'Unload model'),
        ('GET', '/models/{id}/health', 'No', 'Model health'),
        ('POST', '/models/download', 'Admin', 'Download model'),
        ('GET', '/models/download/{task}', 'No', 'Download progress'),
        # Keys
        ('GET', '/keys/my-key', 'JWT', 'View API key'),
        ('POST', '/keys/generate', 'JWT', 'Generate new key'),
        ('GET', '/keys/my-key/stats', 'JWT', 'Key usage stats'),
        ('DELETE', '/keys/my-key', 'JWT', 'Revoke key'),
        ('GET', '/keys/admin/all', 'Admin', 'All keys'),
        ('POST', '/keys/admin/revoke', 'Admin', 'Force revoke'),
        # Quota
        ('GET', '/quota/limits', 'No', 'Default limits'),
        ('GET', '/quota/me', 'JWT', 'My quota'),
        ('PUT', '/quota/admin/user/{roll}', 'Admin', 'Set override'),
        ('GET', '/quota/admin/exceeded', 'Admin', 'Exceeded users'),
        # Guardrails
        ('POST', '/guardrails/check-input', 'JWT', 'Check input'),
        ('POST', '/guardrails/check-output', 'JWT', 'Check output'),
        ('GET', '/guardrails/rules', 'Admin', 'List rules'),
        ('PUT', '/guardrails/rules', 'Admin', 'Update rules'),
        # RAG
        ('POST', '/rag/ingest', 'JWT', 'Upload document'),
        ('GET', '/rag/documents', 'JWT', 'List documents'),
        ('GET', '/rag/documents/{id}', 'JWT', 'Document detail'),
        ('DELETE', '/rag/documents/{id}', 'Admin', 'Delete document'),
        ('POST', '/rag/query', 'JWT', 'RAG-augmented Q&A'),
        ('GET', '/rag/query/{id}/sources', 'JWT', 'Query sources'),
        ('POST', '/rag/collections', 'Admin', 'Create collection'),
        ('GET', '/rag/collections', 'JWT', 'List collections'),
        # Search
        ('POST', '/search/web', 'JWT', 'Web search'),
        ('POST', '/search/wikipedia', 'JWT', 'Wikipedia search'),
        ('POST', '/search/grounded', 'JWT', 'Grounded answer'),
        ('GET', '/search/cache', 'JWT', 'Search cache'),
        # Integration
        ('GET', '/integration/routing-rules', 'No', 'Routing config'),
        ('PUT', '/integration/routing-rules', 'Admin', 'Update routing'),
        ('GET', '/integration/workers', 'No', 'List workers'),
        ('GET', '/integration/workers/{id}', 'No', 'Worker detail'),
        ('POST', '/integration/workers/{id}/drain', 'Admin', 'Drain worker'),
        ('GET', '/integration/queue', 'No', 'Queue status'),
    ]

    styled_table(doc, ['Method', 'Path', 'Auth', 'Description'],
                 all_endpoints, col_widths=[0.6, 2.8, 0.7, 2.5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 12. PROJECT STRUCTURE
    # ══════════════════════════════════════════════════════════
    heading(doc, '12. Project Structure', level=1)

    body(doc, 'The codebase follows a clean, modular layout. Each concern (routing, business logic, data models, schemas) '
         'is separated into its own directory with one file per domain.')

    tree = (
        'MAC/\n'
        '├── mac/                          # Main application package\n'
        '│   ├── __init__.py\n'
        '│   ├── main.py                   # FastAPI entry point, lifespan, seed\n'
        '│   ├── config.py                 # Pydantic settings from .env\n'
        '│   ├── database.py               # Async SQLAlchemy engine & sessions\n'
        '│   ├── routers/                   # 11 API router modules\n'
        '│   │   ├── auth.py               # 19 endpoints\n'
        '│   │   ├── explore.py            # 5 endpoints\n'
        '│   │   ├── query.py              # 6 endpoints\n'
        '│   │   ├── usage.py              # 6 endpoints\n'
        '│   │   ├── models.py             # 7 endpoints\n'
        '│   │   ├── keys.py               # 6 endpoints\n'
        '│   │   ├── quota.py              # 4 endpoints\n'
        '│   │   ├── guardrails.py         # 4 endpoints\n'
        '│   │   ├── rag.py                # 8 endpoints\n'
        '│   │   ├── search.py             # 4 endpoints\n'
        '│   │   └── integration.py        # 6 endpoints\n'
        '│   ├── services/                  # 7 business logic modules\n'
        '│   │   ├── auth_service.py       # User auth, tokens, bcrypt\n'
        '│   │   ├── llm_service.py        # LLM proxy, smart routing\n'
        '│   │   ├── model_service.py      # Model lifecycle management\n'
        '│   │   ├── guardrail_service.py  # Content filtering logic\n'
        '│   │   ├── rag_service.py        # Chunking, embedding, Qdrant\n'
        '│   │   ├── search_service.py     # SearXNG, Wikipedia, grounded\n'
        '│   │   └── usage_service.py      # Usage aggregation queries\n'
        '│   ├── models/                    # SQLAlchemy ORM models\n'
        '│   │   ├── user.py               # User, StudentRegistry, RefreshToken, UsageLog\n'
        '│   │   ├── guardrail.py          # GuardrailRule\n'
        '│   │   ├── quota.py              # QuotaOverride\n'
        '│   │   └── rag.py                # RAGCollection, RAGDocument\n'
        '│   ├── schemas/                   # 11 Pydantic schema files\n'
        '│   ├── middleware/                # Auth + rate limiting\n'
        '│   └── utils/                     # Security helpers\n'
        '├── frontend/                      # PWA frontend\n'
        '│   ├── index.html                # HTML shell\n'
        '│   ├── app.js                    # ~750 lines SPA\n'
        '│   ├── style.css                 # Full theme\n'
        '│   ├── manifest.json             # PWA manifest\n'
        '│   └── sw.js                     # Service worker\n'
        '├── tests/                         # 81+ test cases\n'
        '│   ├── conftest.py               # Async fixtures\n'
        '│   └── test_*.py                 # 11 test files\n'
        '├── docs/                          # Documentation\n'
        '│   ├── Phase-0 through Phase-8   # Per-phase DOCX docs\n'
        '│   └── MAC-API-Design-Document   # Full API spec\n'
        '├── litellm/config.yaml            # Model routing config\n'
        '├── nginx/nginx.conf               # Reverse proxy config\n'
        '├── alembic/                       # Database migrations\n'
        '├── docker-compose.yml             # 7-service orchestration\n'
        '├── Dockerfile                     # Python 3.11 container\n'
        '├── requirements.txt               # 14 Python dependencies\n'
        '└── pytest.ini                     # Test configuration'
    )
    code_block(doc, tree, size=7)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 13. BUILD ROADMAP
    # ══════════════════════════════════════════════════════════
    heading(doc, '13. Build Roadmap & Phase Summary', level=1)

    body(doc, 'The project was built in 8 sequential phases, each delivering a working increment:')

    add_diagram(doc, dia_roadmap())
    body(doc, 'Figure 12: Eight-phase build roadmap showing dependencies.', italic=True, size=10, justify=False)

    styled_table(doc, ['Phase', 'Name', 'Deliverables', 'Depends On'],
    [
        ['1', 'API Endpoints', '35 core REST endpoints — auth, explore, query, usage', 'None'],
        ['2', 'LLM Models', '5 specialist models selected, VRAM planned, Ollama/vLLM configured', 'Phase 1'],
        ['3', 'API–Model Integration', 'LiteLLM proxy wiring, smart routing, health monitoring', 'Phase 1, 2'],
        ['4', 'Usage Control', 'Rate limiting, token quotas, API key lifecycle, sliding-window algorithm', 'Phase 1'],
        ['5', 'Web Interface', 'PWA frontend — dashboard, chat, settings, admin panel', 'Phase 1, 4'],
        ['6', 'Guardrails', 'Input/output content filtering, PII redaction, custom rules', 'Phase 3'],
        ['7', 'Knowledgebase + RAG', 'Document ingestion, vector embedding, Qdrant storage, retrieval chain', 'Phase 3'],
        ['8', 'Retrieval + Search', 'SearXNG web search, Wikipedia, grounded answers with citations', 'Phase 3, 7'],
    ], col_widths=[0.5, 1.3, 3.2, 1.6])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 14. FUTURE SCOPE
    # ══════════════════════════════════════════════════════════
    heading(doc, '14. Future Scope', level=1)

    body(doc, 'While the current platform is fully functional, several enhancements are planned for future development:')

    future_items = [
        ('Real-time GPU monitoring dashboard:', ' Integrate NVIDIA SMI data to show live GPU temperature, VRAM utilisation, and inference throughput on the admin panel.'),
        ('Model fine-tuning interface:', ' Allow faculty to fine-tune models on domain-specific datasets (e.g., college exam papers, lab manuals) through a web interface.'),
        ('Multi-modal support:', ' Expand beyond text and images to support video analysis and diagram generation.'),
        ('Conversation persistence:', ' Store chat sessions server-side to enable cross-device continuity and admin audit trails.'),
        ('Federation:', ' Connect multiple college instances to share GPU resources and model weights across institutions.'),
        ('Automated model selection:', ' Use ML-based classifiers instead of keyword matching for smarter auto-routing.'),
        ('Certificate-based auth:', ' Support client certificates for zero-trust network environments.'),
        ('Prometheus + Grafana monitoring:', ' Production-grade observability with metrics, alerts, and dashboards.'),
        ('WebSocket support:', ' Replace SSE with WebSocket for bidirectional real-time communication.'),
        ('Mobile-native app:', ' Dedicated iOS/Android apps with push notifications for admin alerts.'),
    ]
    for bold_part, rest in future_items:
        bullet(doc, rest, bold_prefix=bold_part)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 15. CONCLUSION
    # ══════════════════════════════════════════════════════════
    heading(doc, '15. Conclusion', level=1)

    body(doc, 'MAC (MBM AI Cloud) demonstrates that a production-quality, self-hosted AI inference platform can be built '
         'entirely with open-source technologies and existing college hardware. The platform provides 500+ students with '
         'free, private, and high-performance access to state-of-the-art language models — for code generation, mathematical '
         'reasoning, general text tasks, image understanding, and speech transcription.')

    body(doc, 'The architecture is deliberately simple and scalable. A modular FastAPI gateway handles authentication, '
         'rate limiting, and content safety for 50+ API endpoints across 11 modules. The LLM inference is decoupled '
         'through LiteLLM, allowing the system to scale from a single lab PC to 30+ GPU nodes by simply registering '
         'new vLLM workers — zero code changes required.')

    body(doc, 'The unified DOB-based authentication flow eliminates the friction of traditional signup/login processes, '
         'while the PWA frontend provides a ChatGPT-like experience that students can install on any device. Comprehensive '
         'guardrails protect against prompt injection, harmful content, and PII exposure. The RAG pipeline enables '
         'faculty to upload textbooks for AI-augmented learning, and the grounded search module provides cited, factual answers.')

    body(doc, 'With 81+ automated tests, Docker Compose deployment, and thorough documentation, the platform is ready '
         'for production deployment at MBM Engineering College and serves as a replicable template for any educational '
         'institution seeking to democratise AI access.')

    # ── Footer spacer ──
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('— End of Report —')
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)
    r.italic = True

    # ── Save ──────────────────────────────────────────────
    out_path = OUT / 'MAC-Final-Project-Report.docx'
    doc.save(str(out_path))
    size_kb = out_path.stat().st_size // 1024
    print(f'\n  ✓ Report generated: {out_path}')
    print(f'  ✓ Size: {size_kb} KB')
    print(f'  ✓ Sections: 15 chapters')
    print(f'  ✓ Diagrams: 12 professional figures')
    print(f'  ✓ Tables: 20+ styled tables')


if __name__ == '__main__':
    generate_report()
